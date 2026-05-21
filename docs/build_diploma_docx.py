import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path("/Users/cukw/FinalWork")
DOCS = ROOT / "docs"
SRC_MD = DOCS / "Диплом_Мониторинг_активности_пользователей_2026.md"
OUT_DOCX = DOCS / "Диплом.docx"
MATERIALS = DOCS / "appendix_materials"
MATERIALS.mkdir(parents=True, exist_ok=True)
SCREENSHOTS = MATERIALS / "screenshots"
SCREENSHOTS.mkdir(parents=True, exist_ok=True)


def run(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True)


def safe_read(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def build_architecture_diagram() -> Path:
    dot_file = MATERIALS / "architecture.dot"
    png_file = MATERIALS / "architecture.png"
    dot = r"""
digraph G {
  rankdir=LR;
  graph [fontname="Times New Roman", fontsize=12, splines=true, overlap=false];
  node  [shape=box, style="rounded,filled", fillcolor="#F9FBFF", color="#2F5D9A", fontname="Times New Roman", fontsize=11];
  edge  [color="#4A4A4A", fontname="Times New Roman", fontsize=10];

  FE [label="Frontend"];
  GW [label="API Gateway"];
  ACT [label="ActivityService"];
  AUTH [label="AuthService"];
  USER [label="UserService"];
  MET [label="MetricsService"];
  NOTIF [label="NotificationService"];
  REP [label="ReportService"];
  AGM [label="AgentManagementService"];
  AGENT [label="LocalEndpointAgent"];
  MQ [label="RabbitMQ", shape=ellipse, fillcolor="#FFF8E5", color="#AA7A00"];

  DBACT [label="DB Activity"];
  DBAUTH [label="DB Auth"];
  DBUSER [label="DB User"];
  DBMET [label="DB Metrics"];
  DBNOTIF [label="DB Notification"];
  DBREP [label="DB Report"];
  DBAGM [label="DB Agent"];

  FE -> GW;
  GW -> ACT;
  GW -> AUTH;
  GW -> USER;
  GW -> MET;
  GW -> NOTIF;
  GW -> REP;
  GW -> AGM;

  AGENT -> ACT [label="CreateActivity"];
  AGENT -> AGM [label="policy / commands / heartbeat"];

  ACT -> MQ [label="domain events"];
  MQ -> MET;
  MQ -> NOTIF;
  MQ -> REP;

  ACT -> DBACT;
  AUTH -> DBAUTH;
  USER -> DBUSER;
  MET -> DBMET;
  NOTIF -> DBNOTIF;
  REP -> DBREP;
  AGM -> DBAGM;
}
""".strip()
    dot_file.write_text(dot + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/dot", "-Tpng", str(dot_file), "-o", str(png_file)])
    return png_file


def build_event_flow_diagram() -> Path:
    puml_file = MATERIALS / "event_flow.puml"
    png_file = MATERIALS / "event_flow.png"
    puml = r"""
@startuml
skinparam defaultFontName "Times New Roman"
skinparam sequenceMessageAlign center
actor "Endpoint Agent" as Agent
participant "ActivityService" as Act
database "Activity DB" as DbA
queue "RabbitMQ" as MQ
participant "MetricsService" as Met
participant "NotificationService" as Notif
participant "ReportService" as Rep

Agent -> Act: CreateActivity
Act -> DbA: Save activity
Act -> Act: Run anomaly rules
Act -> DbA: Save anomaly
Act -> DbA: Save outbox record
Act -> MQ: Publish event

MQ -> Met: Update rollups
MQ -> Notif: Create notification
MQ -> Rep: Update report projection
@enduml
""".strip()
    puml_file.write_text(puml + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/plantuml", "-tpng", str(puml_file), "-o", str(MATERIALS)])
    return png_file


def build_screen_map_diagram() -> Path:
    dot_file = MATERIALS / "screen_map.dot"
    png_file = MATERIALS / "screen_map.png"
    dot = r"""
digraph Screens {
  rankdir=TB;
  graph [fontname="Times New Roman", fontsize=12, splines=ortho];
  node [shape=box, style="rounded,filled", fillcolor="#FFFDF6", color="#8E6F1D", fontname="Times New Roman", fontsize=11];
  edge [color="#5E5E5E", arrowsize=0.8];

  Login [label="Логин\n/login"];
  Dashboard [label="Dashboard\n/dashboard"];
  Agents [label="Agents\n/agents"];
  Analytics [label="Analytics\n/analytics"];
  Reports [label="Reports\n/reports"];
  Settings [label="Settings\n/settings"];
  Users [label="Users\n/users"];

  Login -> Dashboard;
  Dashboard -> Agents;
  Dashboard -> Analytics;
  Dashboard -> Reports;
  Dashboard -> Settings;
  Dashboard -> Users;
}
""".strip()
    dot_file.write_text(dot + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/dot", "-Tpng", str(dot_file), "-o", str(png_file)])
    return png_file


def build_data_model_diagram() -> Path:
    dot_file = MATERIALS / "data_model.dot"
    png_file = MATERIALS / "data_model.png"
    dot = r"""
digraph DataModel {
  rankdir=LR;
  graph [fontname="Times New Roman", fontsize=12, splines=true];
  node [shape=record, fontname="Times New Roman", fontsize=10, color="#2F5D9A"];
  edge [color="#4A4A4A"];

  activities [label="{activities|id\lcomputer_id\ltimestamp\lactivity_type\lrisk_score\l}"];
  anomalies [label="{anomalies|id\lactivity_id\ltype\ldetected_at\l}"];
  outbox [label="{activity_outbox|id\levent_type\lactivity_id\lattempt_count\l}"];
  notifications [label="{notifications|id\luser_id\ltype\ldelivery_status\l}"];
  inbox [label="{processed_event_inbox|id\lconsumer\levent_key\lprocessed_at\l}"];
  agents [label="{agents|id\lcomputer_id\lversion\lstatus\l}"];
  policies [label="{agent_policies|id\lagent_id\lpolicy_version\lupdated_at\l}"];
  commands [label="{agent_commands|id\lagent_id\lcommand_key\lstatus\l}"];

  anomalies -> activities [label="activity_id"];
  outbox -> activities [label="activity_id"];
  policies -> agents [label="agent_id"];
  commands -> agents [label="agent_id"];
}
""".strip()
    dot_file.write_text(dot + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/dot", "-Tpng", str(dot_file), "-o", str(png_file)])
    return png_file


def build_control_plane_diagram() -> Path:
    puml_file = MATERIALS / "control_plane.puml"
    png_file = MATERIALS / "control_plane.png"
    puml = r"""
@startuml
skinparam defaultFontName "Times New Roman"
rectangle "Администратор" as Admin
rectangle "Frontend" as FE
rectangle "Gateway" as GW
rectangle "AgentManagementService" as AGM
rectangle "Endpoint Agent" as Agent

Admin --> FE : настройка политик
FE --> GW : REST API
GW --> AGM : gRPC
AGM --> Agent : policy / command
Agent --> AGM : heartbeat / ack
AGM --> FE : state / results
@enduml
""".strip()
    puml_file.write_text(puml + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/plantuml", "-tpng", str(puml_file), "-o", str(MATERIALS)])
    return png_file


def build_deployment_diagram() -> Path:
    dot_file = MATERIALS / "deployment.dot"
    png_file = MATERIALS / "deployment.png"
    dot = r"""
digraph Deployment {
  rankdir=LR;
  graph [fontname="Times New Roman", fontsize=12, splines=ortho];
  node [shape=box, style="rounded,filled", fontname="Times New Roman", fontsize=10];
  edge [color="#4A4A4A"];

  client [label="Пользовательский браузер", fillcolor="#F4F9FF", color="#2563EB"];
  frontend [label="Frontend\nReact + Nginx\n3000 / 3443", fillcolor="#F4F9FF", color="#2563EB"];
  gateway [label="Gateway\nASP.NET Core\n8080", fillcolor="#FFF8E5", color="#AA7A00"];
  rabbit [label="RabbitMQ\n15672", fillcolor="#FFF8E5", color="#AA7A00"];
  prometheus [label="Prometheus\n9090", fillcolor="#F3FFF6", color="#228B4E"];
  grafana [label="Grafana\n3001", fillcolor="#F3FFF6", color="#228B4E"];

  subgraph cluster_services {
    label="Backend services";
    color="#D6E4FF";
    style="rounded";
    activity [label="ActivityService\n5001 / 5002", fillcolor="#FFFFFF", color="#2F5D9A"];
    auth [label="AuthService\n5003 / 5007", fillcolor="#FFFFFF", color="#2F5D9A"];
    user [label="UserService\n5004 / 5005", fillcolor="#FFFFFF", color="#2F5D9A"];
    metrics [label="MetricsService\n5010 / 5011", fillcolor="#FFFFFF", color="#2F5D9A"];
    notification [label="NotificationService\n5012 / 5017", fillcolor="#FFFFFF", color="#2F5D9A"];
    report [label="ReportService\n5013 / 5014", fillcolor="#FFFFFF", color="#2F5D9A"];
    agentmgmt [label="AgentManagementService\n5015 / 5016", fillcolor="#FFFFFF", color="#2F5D9A"];
  }

  subgraph cluster_db {
    label="PostgreSQL";
    color="#FFE5E5";
    style="rounded";
    db1 [label="activities", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
    db2 [label="auth", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
    db3 [label="users", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
    db4 [label="metrics", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
    db5 [label="notifications", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
    db6 [label="reports", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
    db7 [label="agents", shape=cylinder, fillcolor="#FFFFFF", color="#A14B4B"];
  }

  client -> frontend;
  frontend -> gateway;
  gateway -> activity;
  gateway -> auth;
  gateway -> user;
  gateway -> metrics;
  gateway -> notification;
  gateway -> report;
  gateway -> agentmgmt;
  activity -> rabbit;
  rabbit -> metrics;
  rabbit -> notification;
  rabbit -> report;
  activity -> db1;
  auth -> db2;
  user -> db3;
  metrics -> db4;
  notification -> db5;
  report -> db6;
  agentmgmt -> db7;
  prometheus -> gateway;
  prometheus -> activity;
  prometheus -> auth;
  prometheus -> notification;
  grafana -> prometheus;
}
""".strip()
    dot_file.write_text(dot + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/dot", "-Tpng", str(dot_file), "-o", str(png_file)])
    return png_file


def build_gateway_api_diagram(controller_names: list[str]) -> Path:
    dot_file = MATERIALS / "gateway_api.dot"
    png_file = MATERIALS / "gateway_api.png"
    node_defs = []
    edges = []
    for name in controller_names:
        node_id = re.sub(r"[^A-Za-z0-9_]", "_", name.lower())
        node_defs.append(
            f'  {node_id} [label="{name}Controller", fillcolor="#FFFFFF", color="#2F5D9A"];'
        )
        edges.append(f"  gateway -> {node_id};")
    dot = "\n".join(
        [
            "digraph GatewayApi {",
            '  rankdir=TB;',
            '  graph [fontname="Times New Roman", fontsize=12, splines=ortho];',
            '  node [shape=box, style="rounded,filled", fontname="Times New Roman", fontsize=10];',
            '  edge [color="#4A4A4A"];',
            '  gateway [label="Gateway REST API", fillcolor="#FFF8E5", color="#AA7A00"];',
            *node_defs,
            *edges,
            "}",
        ]
    )
    dot_file.write_text(dot + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/dot", "-Tpng", str(dot_file), "-o", str(png_file)])
    return png_file


def build_observability_diagram() -> Path:
    puml_file = MATERIALS / "observability.puml"
    png_file = MATERIALS / "observability.png"
    puml = r"""
@startuml
skinparam defaultFontName "Times New Roman"
rectangle "Gateway" as GW
rectangle "ActivityService" as ACT
rectangle "AuthService" as AUTH
rectangle "NotificationService" as NOTIF
rectangle "Prometheus" as PROM
rectangle "Grafana" as GRAF
rectangle "Alert Rules" as ALERT
actor "Администратор" as ADMIN

PROM --> GW : scrape /metrics
PROM --> ACT : scrape /metrics
PROM --> AUTH : scrape /metrics
PROM --> NOTIF : scrape /metrics
ALERT --> PROM : expressions / thresholds
GRAF --> PROM : dashboard queries
ADMIN --> GRAF : monitoring / analysis
ADMIN --> ALERT : incident response
@enduml
""".strip()
    puml_file.write_text(puml + "\n", encoding="utf-8")
    run(["/opt/homebrew/bin/plantuml", "-tpng", str(puml_file), "-o", str(MATERIALS)])
    return png_file


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def set_section_margins(section) -> None:
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


def set_run_font(run, size: int = 14, bold: bool = False, font_name: str = "Times New Roman") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def format_paragraph(paragraph, justify: bool = True, first_indent: float = 1.25, line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(first_indent)
    paragraph.paragraph_format.line_spacing_rule = line_spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT


def add_text_paragraph(doc: Document, text: str, justify: bool = True, first_indent: float = 1.25, size: int = 14) -> None:
    p = doc.add_paragraph()
    format_paragraph(p, justify=justify, first_indent=first_indent)
    r = p.add_run(text)
    set_run_font(r, size=size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=14, bold=False)
    for rr in p.runs:
        set_run_font(rr, size=14, bold=False)
    spacer = doc.add_paragraph()
    format_paragraph(spacer, justify=False, first_indent=0)


def add_unlisted_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=14, bold=False)
    spacer = doc.add_paragraph()
    format_paragraph(spacer, justify=False, first_indent=0)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=12)


def add_centered_picture(doc: Document, image_path: Path, width: float, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.first_line_indent = Cm(0)
    rr = cp.add_run(caption)
    set_run_font(rr, size=12)


def add_code_block(doc: Document, title: str, code: str) -> None:
    add_text_paragraph(doc, title)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(code.strip())
    set_run_font(r, size=10, font_name="Courier New")


def add_table_font(table, size: int = 12) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(r, size=size)


def format_appendix_content(doc: Document) -> None:
    in_appendix = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Приложения" and paragraph.style.name.startswith("Heading"):
            in_appendix = True

        if not in_appendix:
            continue

        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        if not paragraph.style.name.startswith("Heading"):
            for run in paragraph.runs:
                run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    for run in paragraph.runs:
                        run.font.size = Pt(12)


def add_page_number(footer_paragraph) -> None:
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = footer_paragraph.add_run()
    set_run_font(run, size=14)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def parse_main_blocks(text: str) -> list[str]:
    normalized = text.replace("\r\n", "\n")
    blocks = [block.strip() for block in re.split(r"\n\s*\n", normalized) if block.strip()]
    if blocks and blocks[0].startswith("Тема дипломной работы:"):
        blocks = blocks[1:]
    return blocks


def slice_between(text: str, start: str, end: str) -> str:
    start_idx = text.find(start)
    if start_idx == -1:
        return ""
    end_idx = text.find(end, start_idx)
    if end_idx == -1:
        end_idx = len(text)
    return text[start_idx:end_idx].strip()


def extract_context(path: Path, marker: str, before: int = 3, after: int = 10) -> str:
    lines = safe_read(path).splitlines()
    for idx, line in enumerate(lines):
        if marker in line:
            start = max(0, idx - before)
            end = min(len(lines), idx + after)
            return "\n".join(lines[start:end]).strip()
    return "\n".join(lines[: min(len(lines), 20)]).strip()


def extract_compose_services(path: Path) -> list[tuple[str, str, str]]:
    purpose_map = {
        "frontend": "Web-клиент и обратный прокси для пользовательского доступа",
        "gateway": "Единая REST-точка входа и маршрутизация запросов к gRPC-сервисам",
        "activityservice": "Прием активности, расчет рисков и публикация доменных событий",
        "authservice": "Аутентификация, выпуск токенов и сопровождение сессий",
        "userservice": "Учетные записи, компьютеры и организационные атрибуты",
        "metricservice": "Агрегация операционных метрик и аналитических срезов",
        "notificationservice": "Формирование и доставка уведомлений, retry и DLQ",
        "reportservice": "Построение отчетных представлений и выгрузок",
        "agentmanagementservice": "Control plane для политик, heartbeat и команд агентов",
        "activityagent": "Сервисный генератор активности для стендов и smoke-проверок",
        "rabbitmq": "Транспорт асинхронных событий и межсервисной интеграции",
        "prometheus": "Сбор технических метрик и вычисление alert-правил",
        "grafana": "Визуализация дашбордов и наблюдаемость платформы",
        "postgres-activity": "Хранилище ActivityService",
        "postgres-auth": "Хранилище AuthService",
        "postgres-user": "Хранилище UserService",
        "postgres-metrics": "Хранилище MetricsService",
        "postgres-notification": "Хранилище NotificationService",
        "postgres-report": "Хранилище ReportService",
        "postgres-agent": "Хранилище AgentManagementService",
    }

    rows: list[tuple[str, str, str]] = []
    current_name = ""
    current_ports: list[str] = []
    in_ports = False

    def flush() -> None:
        nonlocal current_name, current_ports
        if current_name:
            rows.append(
                (
                    current_name,
                    ", ".join(current_ports) if current_ports else "без публикации наружу",
                    purpose_map.get(current_name, "Сервис инфраструктурного или прикладного уровня"),
                )
            )
        current_name = ""
        current_ports = []

    for line in safe_read(path).splitlines():
        service_match = re.match(r"^  ([a-z0-9-]+):\s*$", line)
        if service_match:
            flush()
            current_name = service_match.group(1)
            in_ports = False
            continue

        if not current_name:
            continue

        if re.match(r"^    ports:\s*$", line):
            in_ports = True
            continue

        if in_ports:
            port_match = re.match(r'^      - "([^"]+)"', line)
            if port_match:
                current_ports.append(port_match.group(1))
                continue
            if re.match(r"^    [A-Za-z0-9_-]+:", line):
                in_ports = False

    flush()
    return rows


def join_route(base_route: str, suffix: str) -> str:
    base = "/" + base_route.strip("/")
    suffix = suffix.strip()
    if not suffix:
        return base
    return f"{base}/{suffix.lstrip('/')}"


def extract_gateway_routes(controller_dir: Path) -> list[tuple[str, str, str]]:
    rows: list[tuple[str, str, str]] = []
    for path in sorted(controller_dir.glob("*Controller.cs")):
        controller = path.stem.replace("Controller", "")
        base_route = ""
        for line in safe_read(path).splitlines():
            route_match = re.search(r'\[Route\("([^"]+)"\)\]', line)
            if route_match:
                base_route = route_match.group(1)

            http_match = re.search(r'\[(HttpGet|HttpPost|HttpPut|HttpDelete|HttpPatch)(?:\("([^"]*)"\))?\]', line)
            if http_match and base_route:
                verb = http_match.group(1).replace("Http", "").upper()
                suffix = http_match.group(2) or ""
                rows.append((controller, verb, join_route(base_route, suffix)))
    return rows


def extract_prometheus_jobs(path: Path) -> list[tuple[str, str]]:
    jobs: list[tuple[str, str]] = []
    current_job = ""
    for line in safe_read(path).splitlines():
        job_match = re.search(r"job_name:\s*([A-Za-z0-9_-]+)", line)
        if job_match:
            current_job = job_match.group(1)
            continue
        target_match = re.search(r'targets:\s*\["([^"]+)"\]', line)
        if target_match and current_job:
            jobs.append((current_job, target_match.group(1)))
            current_job = ""
    return jobs


def extract_alert_rules(path: Path) -> list[tuple[str, str, str]]:
    rules: list[tuple[str, str, str]] = []
    current_name = ""
    current_for = ""
    for line in safe_read(path).splitlines():
        alert_match = re.search(r"alert:\s*([A-Za-z0-9_]+)", line)
        if alert_match:
            current_name = alert_match.group(1)
            current_for = ""
            continue
        for_match = re.search(r"for:\s*([0-9a-zA-Z]+)", line)
        if for_match and current_name:
            current_for = for_match.group(1)
            continue
        severity_match = re.search(r"severity:\s*([A-Za-z0-9_-]+)", line)
        if severity_match and current_name:
            rules.append((current_name, severity_match.group(1), current_for or "не указан"))
            current_name = ""
            current_for = ""
    return rules


def build_document() -> None:
    src_text = safe_read(SRC_MD)
    blocks = parse_main_blocks(src_text)
    compose_services = extract_compose_services(ROOT / "docker-compose.yml")
    gateway_routes = extract_gateway_routes(ROOT / "Backend" / "gateway" / "src" / "Controllers")
    prometheus_jobs = extract_prometheus_jobs(ROOT / "observability" / "prometheus" / "prometheus.yml")
    alert_rules = extract_alert_rules(ROOT / "observability" / "prometheus" / "backend-alerts.yml")

    architecture_img = build_architecture_diagram()
    event_flow_img = build_event_flow_diagram()
    screen_map_img = build_screen_map_diagram()
    data_model_img = build_data_model_diagram()
    control_plane_img = build_control_plane_diagram()
    deployment_img = build_deployment_diagram()
    gateway_api_img = build_gateway_api_diagram(sorted({row[0] for row in gateway_routes}))
    observability_img = build_observability_diagram()

    app_routes = slice_between(safe_read(ROOT / "Frontend" / "src" / "App.js"), "return (\n    <Routes>", "function App()")
    activity_sql = safe_read(ROOT / "Backend" / "services" / "ActivityService" / "db" / "initActivity.sql")
    agent_sql = safe_read(ROOT / "Backend" / "services" / "AgentManagementService" / "db" / "InitAgent.sql")
    notification_sql = safe_read(ROOT / "Backend" / "services" / "NotificationService" / "db" / "initNotificationService.sql")
    compose_snippet = safe_read(ROOT / "docker-compose.yml")
    prometheus_snippet = safe_read(ROOT / "observability" / "prometheus" / "prometheus.yml")
    e2e_snippet = safe_read(ROOT / "scripts" / "e2e_smoke.sh")
    ci_snippet = safe_read(ROOT / ".github" / "workflows" / "ci.yml")
    agent_proto = safe_read(ROOT / "Backend" / "gateway" / "src" / "Protos" / "agent.proto")
    controller_snippet = extract_context(
        ROOT / "Backend" / "gateway" / "src" / "Controllers" / "AgentController.cs",
        '[HttpPost("agents/{id:long}/commands")]',
        before=4,
        after=18,
    )
    activity_service_snippet = extract_context(
        ROOT / "Backend" / "services" / "ActivityService" / "Services" / "Models" / "ActivityServiceImpl.cs",
        "OutboxEventEnvelopeFactory.CreateActivityCreated",
        before=6,
        after=14,
    )
    runner_snippet = extract_context(
        ROOT / "LocalEndpointAgent" / "src" / "endpoint_agent" / "runner.py",
        "async def",
        before=0,
        after=20,
    )

    doc = Document()
    enable_update_fields_on_open(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = False

    set_section_margins(doc.sections[0])

    title_lines = [
        "ДЕПАРТАМЕНТ ОБРАЗОВАНИЯ И НАУКИ ГОРОДА МОСКВЫ",
        "Государственное бюджетное профессиональное образовательное учреждение города Москвы",
        "«Колледж связи №54 имени П.М. Вострухина»",
        "",
        "ДОПУСКАЮ К ЗАЩИТЕ",
        "_______________________",
        "Заместитель директора по УПР ____________________",
        "_________________",
        "(дата)",
        "",
        "ДИПЛОМНАЯ РАБОТА",
        "",
        "Мониторинг активности пользователей",
        "(тема)",
        "выполнена",
        "студентом группы",
        "________________",
        "(номер группы)",
        "________________________________",
        "(И. О. Фамилия)",
        "(подпись, дата)",
        "Основная профессиональная образовательная программа по специальности",
        "09.02.07 Информационные системы и программирование",
        "(шифр и наименование специальности)",
        "Форма обучения",
        "очная",
        "Руководитель",
        "преподаватель",
        "________________________________",
        "(ученая степень, должность, И. О. Фамилия)",
        "(подпись, дата)",
        "Председатель предметной (междисциплинарной, модульной) комиссии",
        "________________________________",
        "(И. О. Фамилия)",
        "(подпись, дата)",
        "Москва",
        "2026",
    ]

    for line in title_lines:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line in {
            "ДОПУСКАЮ К ЗАЩИТЕ",
            "Заместитель директора по УПР ____________________",
            "_________________",
            "(дата)",
        }:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        set_run_font(r, size=14)

    doc.add_page_break()
    for line in [
        "Государственное бюджетное профессиональное образовательное учреждение города Москвы",
        "«Колледж связи №54 имени П.М. Вострухина»",
        "",
        "УТВЕРЖДАЮ",
        "_______________________",
        "Заместитель директора по УПР ____________________",
        "_________________",
        "(дата)",
        "",
        "ЗАДАНИЕ НА ДИПЛОМНУЮ РАБОТУ",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if "ЗАДАНИЕ" in line else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        set_run_font(r, size=14)

    assignment_rows = [
        "Студенту ________________________________________________",
        "I. Тема дипломной работы: Мониторинг активности пользователей.",
        "II. Срок сдачи студентом законченной работы: ______________________________.",
        "III. Исходные данные: исходный код и конфигурации репозитория FinalWork, материалы преддипломной практики, методические указания колледжа.",
        "IV. Перечень подлежащих разработке вопросов:",
        "1. Анализ предметной области автоматизации и определение цели, задач и требований к информационной системе.",
        "2. Выполнение системного анализа объекта исследования, обзор существующих решений и подготовка проекта управления разработкой.",
        "3. Разработка технического задания, проектирование информационной системы и проектирование базы данных.",
        "4. Реализация физической модели базы данных, интерфейса и функционала приложения.",
        "5. Тестирование приложения, проверка демонстрационного стенда и подготовка руководства пользователя.",
        "6. Подготовка приложений со схемами, картой экранов, листингами, материалами тестирования, руководством пользователя и скриншотами.",
        "V. Перечень приложений:",
        "1. Техническое задание.",
        "2. Карта экранов и маршрутов web-приложения.",
        "3. Листинг базы данных.",
        "4. Листинг приложения.",
        "5. Схема развертывания и инфраструктуры.",
        "6. Карта API и сервисных контрактов.",
        "7. Материалы тестирования и наблюдаемости.",
        "8. Руководство пользователя.",
        "9. Скриншоты демонстрационного стенда.",
        "VI. Дата выдачи задания: ______________________________.",
        "Руководитель ______________________________",
        "Задание принял к исполнению ______________________________",
    ]
    for row in assignment_rows:
        add_text_paragraph(doc, row)

    doc.add_page_break()
    add_unlisted_heading(doc, "КАЛЕНДАРНЫЙ ПЛАН ДИПЛОМНОЙ РАБОТЫ")
    plan = doc.add_table(rows=1, cols=4)
    plan.style = "Table Grid"
    plan.rows[0].cells[0].text = "Этап"
    plan.rows[0].cells[1].text = "Содержание этапа"
    plan.rows[0].cells[2].text = "Срок"
    plan.rows[0].cells[3].text = "Отметка"
    plan_rows = [
        ("1", "Анализ предметной области автоматизации и постановка цели и задач системы", "__________", "__________"),
        ("2", "Системный анализ объекта исследования, требований и существующих решений", "__________", "__________"),
        ("3", "Проектирование архитектуры системы и модели базы данных", "__________", "__________"),
        ("4", "Разработка функционала приложения, интерфейса и физической модели БД", "__________", "__________"),
        ("5", "Тестирование приложения, проверка демонстрационного стенда и фиксация скриншотов", "__________", "__________"),
        ("6", "Подготовка руководства пользователя, приложений и итогового оформления", "__________", "__________"),
    ]
    for row in plan_rows:
        cells = plan.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    add_table_font(plan)

    doc.add_page_break()
    add_unlisted_heading(doc, "Оглавление")
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    add_field(toc, ' TOC \\o "1-3" \\h \\z \\u ')

    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_margins(sec2)
    sec2.footer.is_linked_to_previous = False
    sect_pr = sec2._sectPr
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), "5")
    sect_pr.append(pg_num)
    add_page_number(sec2.footer.paragraphs[0])

    major_headings = {
        "ВВЕДЕНИЕ",
        "Введение",
        "Глава 1. Анализ и проектирование предметной области",
        "Глава 1. Аналитическая часть",
        "Глава I. Теоретическая часть",
        "Глава 2. Разработка информационной системы",
        "Глава 2. Практическая часть",
        "Глава II. Практическая часть",
        "Заключение",
        "Список использованных источников",
        "Список литературы",
        "Приложения",
    }

    in_refs = False
    for block in blocks:
        if block in major_headings:
            if block == "Приложения":
                break
            if block != "Введение":
                doc.add_page_break()
            add_heading(doc, block.title() if block == "ВВЕДЕНИЕ" else block, level=1)
            in_refs = block in {"Список использованных источников", "Список литературы"}
            continue

        if re.match(r"^\d+\.\d+\.?\s", block):
            add_heading(doc, block, level=2)
            continue

        if block in {
            "Вывод по теоретической части",
            "Вывод по практической части",
            "Выводы по практической части",
        }:
            add_heading(doc, block, level=2)
            continue

        if in_refs and re.match(r"^\[\d+\]", block):
            add_text_paragraph(doc, block, justify=False)
            continue

        add_text_paragraph(doc, block)

    doc.add_section(WD_SECTION.NEW_PAGE)
    appendix_section = doc.sections[-1]
    set_section_margins(appendix_section)
    appendix_section.footer.is_linked_to_previous = False
    add_heading(doc, "Приложения", level=1)

    doc.add_page_break()
    add_heading(doc, "Приложение 1. Техническое задание", level=2)
    add_text_paragraph(
        doc,
        "Назначение разработки состоит в создании корпоративной системы мониторинга активности пользователей, обеспечивающей сбор событий с рабочих станций, выявление аномального поведения, формирование уведомлений и подготовку аналитической отчетности."
    )
    add_text_paragraph(
        doc,
        "Объект автоматизации представляет собой цифровую среду организации, в рамках которой требуется централизованно контролировать действия пользователей, состояние конечных устройств и выполнение политик безопасности."
    )
    add_table_caption(doc, "Таблица 1.1. Основные параметры технического задания")
    tz_table = doc.add_table(rows=1, cols=2)
    tz_table.style = "Table Grid"
    tz_table.rows[0].cells[0].text = "Параметр"
    tz_table.rows[0].cells[1].text = "Содержание"
    tz_rows = [
        ("Категории пользователей", "Администратор системы, оператор мониторинга, endpoint-агент, сервисы backend-контура"),
        ("Основные функции", "Аутентификация, сбор активности, обнаружение аномалий, уведомления, отчеты, управление политиками и командами агентов"),
        ("Нефункциональные требования", "Отказоустойчивость, ролевая авторизация, аудит действий, масштабируемость, изоляция хранилищ, поддержка событийной интеграции"),
        ("Технологическая основа", "React, ASP.NET Core, gRPC, RabbitMQ, PostgreSQL, Python, Docker Compose"),
    ]
    for key, value in tz_rows:
        row = tz_table.add_row().cells
        row[0].text = key
        row[1].text = value
    add_table_font(tz_table)
    add_table_caption(doc, "Таблица 1.2. Требования к входным и выходным данным")
    req_table = doc.add_table(rows=1, cols=2)
    req_table.style = "Table Grid"
    req_table.rows[0].cells[0].text = "Группа требований"
    req_table.rows[0].cells[1].text = "Содержание"
    req_rows = [
        ("Входные данные", "Телеметрия действий пользователя, heartbeat агентов, команды администрирования, настройки политик"),
        ("Выходные данные", "Агрегированные показатели, выявленные аномалии, уведомления, отчеты, журналы аудита"),
        ("Ограничения", "Разграничение доступа по ролям, защита персональных данных, идемпотентность команд и событий"),
        ("Критерии приемки", "Корректная регистрация пользователей, сохранение настроек, управление агентами и воспроизводимость отчетов"),
    ]
    for key, value in req_rows:
        row = req_table.add_row().cells
        row[0].text = key
        row[1].text = value
    add_table_font(req_table)
    add_centered_picture(
        doc,
        architecture_img,
        6.1,
        "Рисунок 1.1. Архитектурная схема программного комплекса мониторинга активности пользователей",
    )
    add_text_paragraph(
        doc,
        "Схема фиксирует состав основных компонентов и показывает, что пользовательский интерфейс, gateway, доменные сервисы, брокер сообщений, базы данных и endpoint-агент образуют единый контур обработки активности."
    )

    doc.add_page_break()
    add_heading(doc, "Приложение 2. Карта экранов", level=2)
    add_text_paragraph(
        doc,
        "Карта экранов построена на основании маршрутов frontend-приложения и отражает основную навигацию административной панели."
    )
    add_centered_picture(
        doc,
        screen_map_img,
        5.8,
        "Рисунок 2.1. Карта экранов и маршрутов web-приложения",
    )
    add_text_paragraph(
        doc,
        "После авторизации пользователь переходит в защищенные разделы приложения. Карта используется как основа для проверки навигации и сопоставления экранов со скриншотами демонстрационного стенда."
    )
    add_table_caption(doc, "Таблица 2.1. Маршруты web-приложения")
    routes_table = doc.add_table(rows=1, cols=3)
    routes_table.style = "Table Grid"
    routes_table.rows[0].cells[0].text = "Маршрут"
    routes_table.rows[0].cells[1].text = "Экран"
    routes_table.rows[0].cells[2].text = "Назначение"
    for route, page, purpose in [
        ("/login", "Login", "Авторизация пользователя"),
        ("/dashboard", "Dashboard", "Оперативная панель и KPI"),
        ("/agents", "Agents", "Управление endpoint-агентами и командами"),
        ("/analytics", "Analytics", "Аналитические срезы и drill-down"),
        ("/reports", "Reports", "Формирование и экспорт отчетов"),
        ("/settings", "Settings", "Настройки безопасности, списков доступа и правил"),
        ("/users", "Users", "Просмотр и управление пользователями"),
    ]:
        row = routes_table.add_row().cells
        row[0].text = route
        row[1].text = page
        row[2].text = purpose
    add_table_font(routes_table)
    add_table_caption(doc, "Таблица 2.2. Роли и основные разделы интерфейса")
    role_table = doc.add_table(rows=1, cols=3)
    role_table.style = "Table Grid"
    role_table.rows[0].cells[0].text = "Роль"
    role_table.rows[0].cells[1].text = "Ключевые разделы"
    role_table.rows[0].cells[2].text = "Практическое назначение"
    for role, sections, result in [
        ("Администратор", "Dashboard, Agents, Settings, Users", "Оперативный контроль платформы и настройка политик"),
        ("Оператор мониторинга", "Dashboard, Analytics, Reports", "Анализ событий, аномалий и отчетных срезов"),
        ("Сервисный контур", "API / фоновые процессы", "Передача данных, обработка событий и обновление представлений"),
    ]:
        row = role_table.add_row().cells
        row[0].text = role
        row[1].text = sections
        row[2].text = result
    add_table_font(role_table)

    doc.add_page_break()
    add_heading(doc, "Приложение 3. Листинг базы данных", level=2)
    add_text_paragraph(
        doc,
        "Ниже приведены фрагменты SQL-структур, описывающих хранение событий активности, управление агентами и контур уведомлений."
    )
    add_centered_picture(
        doc,
        data_model_img,
        6.0,
        "Рисунок 3.1. Укрупненная схема ключевых сущностей системы",
    )
    add_text_paragraph(
        doc,
        "Схема связывает транзакционные сущности активности с агентами, политиками, командами, уведомлениями и служебными таблицами надежной доставки событий."
    )
    add_code_block(doc, "Фрагмент SQL-описания таблиц ActivityService.", "\n".join(activity_sql.splitlines()[:45]))
    add_code_block(doc, "Фрагмент SQL-описания таблиц AgentManagementService.", "\n".join(agent_sql.splitlines()[:60]))
    add_code_block(doc, "Фрагмент SQL-описания таблиц NotificationService.", "\n".join(notification_sql.splitlines()[:45]))
    add_table_caption(doc, "Таблица 3.1. Словарь ключевых сущностей базы данных")
    dictionary_table = doc.add_table(rows=1, cols=3)
    dictionary_table.style = "Table Grid"
    dictionary_table.rows[0].cells[0].text = "Сущность"
    dictionary_table.rows[0].cells[1].text = "Ключевые поля"
    dictionary_table.rows[0].cells[2].text = "Назначение"
    for entity, fields, purpose in [
        ("activities", "id, computer_id, timestamp, activity_type, risk_score", "Фиксация фактов активности и расчетов риска"),
        ("anomalies", "id, activity_id, type, detected_at", "Хранение выявленных отклонений и причин срабатывания"),
        ("activity_outbox", "id, event_type, activity_id, attempt_count", "Гарантированная публикация доменных событий"),
        ("agents", "id, computer_id, version, status", "Сведения об endpoint-агентах и их состоянии"),
        ("agent_commands", "id, agent_id, command_key, status", "Очередь управляющих команд и подтверждений"),
        ("notifications", "id, user_id, type, delivery_status", "Доставка уведомлений пользователям системы"),
    ]:
        row = dictionary_table.add_row().cells
        row[0].text = entity
        row[1].text = fields
        row[2].text = purpose
    add_table_font(dictionary_table)

    doc.add_page_break()
    add_heading(doc, "Приложение 4. Листинг приложения", level=2)
    add_text_paragraph(
        doc,
        "Приложение содержит выборку реальных фрагментов программного кода, отражающих маршрутизацию клиентского приложения, работу серверного API и обработку событий в backend-контуре."
    )
    add_centered_picture(
        doc,
        event_flow_img,
        6.0,
        "Рисунок 4.1. Сценарий обработки события активности в серверной части",
    )
    add_text_paragraph(
        doc,
        "Последовательность показывает, как первичное событие активности превращается в сохраненную запись, событие outbox и несколько независимых проекций для метрик, уведомлений и отчетности."
    )
    add_code_block(doc, "Фрагмент frontend-маршрутизации приложения.", app_routes)
    add_code_block(doc, "Фрагмент контроллера управления агентами.", controller_snippet)
    add_code_block(doc, "Фрагмент логики публикации доменных событий.", activity_service_snippet)
    add_code_block(doc, "Фрагмент управляющего цикла endpoint-агента.", runner_snippet)

    doc.add_page_break()
    add_heading(doc, "Приложение 5. Схема развертывания и инфраструктуры", level=2)
    add_text_paragraph(
        doc,
        "Развертывание платформы описано через docker-compose и включает пользовательский frontend-контур, gateway, прикладные сервисы, брокер сообщений, отдельные базы данных, а также средства наблюдаемости."
    )
    add_centered_picture(
        doc,
        deployment_img,
        6.2,
        "Рисунок 5.1. Укрупненная схема развертывания платформы в контейнерной среде",
    )
    add_text_paragraph(
        doc,
        "Схема развертывания дополняется перечнем контейнеров и опубликованных портов, что позволяет соотнести логическую архитектуру с фактической инфраструктурной конфигурацией."
    )
    add_table_caption(doc, "Таблица 5.1. Сервисы и порты инфраструктурного контура")
    infra_table = doc.add_table(rows=1, cols=3)
    infra_table.style = "Table Grid"
    infra_table.rows[0].cells[0].text = "Сервис"
    infra_table.rows[0].cells[1].text = "Порты"
    infra_table.rows[0].cells[2].text = "Назначение"
    for service, ports, purpose in compose_services:
        row = infra_table.add_row().cells
        row[0].text = service
        row[1].text = ports
        row[2].text = purpose
    add_table_font(infra_table)
    add_code_block(doc, "Фрагмент docker-compose-конфигурации.", "\n".join(compose_snippet.splitlines()[:120]))

    doc.add_page_break()
    add_heading(doc, "Приложение 6. Карта API и сервисных контрактов", level=2)
    add_text_paragraph(
        doc,
        "Данный раздел объединяет REST-маршруты gateway и gRPC-контракты взаимодействия с AgentManagementService, что позволяет отразить как внешний, так и внутренний слой интерфейсов системы."
    )
    add_centered_picture(
        doc,
        gateway_api_img,
        6.2,
        "Рисунок 6.1. Карта основных контроллеров gateway и прикладных подсистем",
    )
    add_text_paragraph(
        doc,
        "Карта API показывает, какие контроллеры доступны через gateway и какие подсистемы участвуют во внешнем REST-контуре приложения."
    )
    add_table_caption(doc, "Таблица 6.1. Основные REST-маршруты gateway")
    api_table = doc.add_table(rows=1, cols=3)
    api_table.style = "Table Grid"
    api_table.rows[0].cells[0].text = "Контроллер"
    api_table.rows[0].cells[1].text = "Метод"
    api_table.rows[0].cells[2].text = "Маршрут"
    for controller, method, route in gateway_routes[:28]:
        row = api_table.add_row().cells
        row[0].text = controller
        row[1].text = method
        row[2].text = route
    add_table_font(api_table)
    add_code_block(doc, "Фрагмент gRPC-контракта управления агентами.", "\n".join(agent_proto.splitlines()[:140]))

    doc.add_page_break()
    add_heading(doc, "Приложение 7. Материалы тестирования и наблюдаемости", level=2)
    add_text_paragraph(
        doc,
        "Тестовый и эксплуатационный контур проекта подтверждается сценариями smoke-проверок, конфигурацией CI-пайплайна, а также средствами мониторинга на базе Prometheus и Grafana."
    )
    add_centered_picture(
        doc,
        observability_img,
        5.8,
        "Рисунок 7.1. Контур сбора технических метрик и визуализации состояния сервисов",
    )
    add_text_paragraph(
        doc,
        "Контур наблюдаемости дополняет функциональную проверку и показывает, как Prometheus, Grafana и alert-правила используются для эксплуатационного контроля сервисов."
    )
    add_table_caption(doc, "Таблица 7.1. Сценарии smoke-проверки")
    smoke_table = doc.add_table(rows=1, cols=2)
    smoke_table.style = "Table Grid"
    smoke_table.rows[0].cells[0].text = "Шаг smoke-проверки"
    smoke_table.rows[0].cells[1].text = "Проверяемый результат"
    for step, result in [
        ("Register and login", "Создание пользователя и получение JWT-токена через gateway"),
        ("CRUD app settings", "Проверка сохранения и чтения системных настроек"),
        ("Create user+computer", "Создание пользователя с обязательной привязкой компьютера"),
        ("Delete user cascade", "Удаление сущности пользователя и связанных записей"),
        ("Agent command endpoints", "Постановка команд block/unblock и чтение истории"),
        ("Logout", "Корректное завершение сессии пользователя"),
        ("Smoke E2E completed", "Подтверждение сквозной работоспособности основных контуров"),
    ]:
        row = smoke_table.add_row().cells
        row[0].text = step
        row[1].text = result
    add_table_font(smoke_table)
    add_table_caption(doc, "Таблица 7.2. Задания Prometheus")
    jobs_table = doc.add_table(rows=1, cols=2)
    jobs_table.style = "Table Grid"
    jobs_table.rows[0].cells[0].text = "Prometheus job"
    jobs_table.rows[0].cells[1].text = "Цель сбора метрик"
    for job, target in prometheus_jobs:
        row = jobs_table.add_row().cells
        row[0].text = job
        row[1].text = target
    add_table_font(jobs_table)
    add_table_caption(doc, "Таблица 7.3. Alert-правила backend-контура")
    alert_table = doc.add_table(rows=1, cols=3)
    alert_table.style = "Table Grid"
    alert_table.rows[0].cells[0].text = "Alert"
    alert_table.rows[0].cells[1].text = "Severity"
    alert_table.rows[0].cells[2].text = "Интервал фиксации"
    for alert_name, severity, duration in alert_rules:
        row = alert_table.add_row().cells
        row[0].text = alert_name
        row[1].text = severity
        row[2].text = duration
    add_table_font(alert_table)
    add_code_block(doc, "Фрагмент smoke E2E-скрипта.", "\n".join(e2e_snippet.splitlines()[:120]))
    add_code_block(doc, "Фрагмент конфигурации CI-пайплайна.", "\n".join(ci_snippet.splitlines()[:80]))
    add_code_block(doc, "Фрагмент конфигурации Prometheus.", "\n".join(prometheus_snippet.splitlines()[:80]))

    doc.add_page_break()
    add_heading(doc, "Приложение 8. Руководство пользователя", level=2)
    add_text_paragraph(
        doc,
        "После запуска системы пользователь проходит процедуру авторизации через страницу входа. При успешной аутентификации выполняется переход на панель Dashboard, где отображаются основные показатели активности, аномалий и состояния агентов."
    )
    add_text_paragraph(
        doc,
        "Для работы с endpoint-агентами используется раздел Agents. В нем доступны просмотр статусов, изменение политик, отправка команд блокировки и разблокировки, а также контроль результатов выполнения команд."
    )
    add_text_paragraph(
        doc,
        "Раздел Analytics предназначен для анализа накопленных данных по периодам и фильтрам. Раздел Reports применяется для получения сводной отчетности и экспорта агрегированных представлений. Раздел Settings используется для настройки списков доступа, правил уведомлений и параметров мониторинга."
    )
    add_text_paragraph(
        doc,
        "Перед началом работы пользователь открывает адрес демонстрационного стенда http://2.26.89.86 и проходит авторизацию. Для демонстрационного доступа используются учетные данные admin / admin123. После успешного входа открывается административная панель, а все основные действия выполняются через боковое меню приложения."
    )
    add_text_paragraph(
        doc,
        "Рекомендуемый порядок работы состоит из нескольких шагов: сначала пользователь проверяет Dashboard и общее состояние системы, затем при необходимости переходит к аналитике, отчетам, агентам, настройкам или пользователям. После завершения работы необходимо выполнить выход из учетной записи, чтобы закрыть административную сессию."
    )
    add_table_caption(doc, "Таблица 8.1. Основные действия пользователя")
    guide_table = doc.add_table(rows=1, cols=2)
    guide_table.style = "Table Grid"
    guide_table.rows[0].cells[0].text = "Действие пользователя"
    guide_table.rows[0].cells[1].text = "Результат"
    guide_rows = [
        ("Открытие http://2.26.89.86", "Переход на демонстрационный стенд через браузер"),
        ("Вход под admin / admin123", "Получение доступа к административной панели стенда"),
        ("Вход в систему через /login", "Проверка учетных данных и открытие административной панели"),
        ("Переход на Dashboard", "Просмотр KPI, последних событий и аномалий"),
        ("Переход на Agents", "Управление списком агентов, политиками и командами"),
        ("Переход на Analytics", "Анализ активности по периодам, категориям и drill-down"),
        ("Переход на Reports", "Формирование периодных отчетов и экспорт результатов"),
        ("Переход на Settings", "Настройка мониторинга, списков доступа и alert-правил"),
        ("Переход на Users", "Просмотр и сопровождение пользователей и связанных компьютеров"),
        ("Выход из системы", "Завершение административной сессии и возврат к странице входа"),
    ]
    for action, result in guide_rows:
        row = guide_table.add_row().cells
        row[0].text = action
        row[1].text = result
    add_table_font(guide_table)
    add_table_caption(doc, "Таблица 8.2. Назначение основных разделов интерфейса")
    sections_table = doc.add_table(rows=1, cols=3)
    sections_table.style = "Table Grid"
    sections_table.rows[0].cells[0].text = "Раздел"
    sections_table.rows[0].cells[1].text = "Что делает пользователь"
    sections_table.rows[0].cells[2].text = "Когда используется"
    for section, action, when in [
        ("Dashboard", "Проверяет KPI, последние события, аномалии и состояние агентов", "В начале смены и при оперативном контроле"),
        ("Analytics", "Выбирает период, фильтрует события и анализирует детализацию активности", "При разборе причин отклонений и подозрительных действий"),
        ("Reports", "Формирует дневные, недельные, месячные и произвольные отчеты", "Для подготовки сводок и передачи результатов руководителю"),
        ("Agents", "Проверяет статусы агентов, политики, команды блокировки и разблокировки", "При администрировании рабочих станций и реагировании на риски"),
        ("Settings", "Настраивает мониторинг, списки доступа, уведомления и правила", "При первичной настройке и изменении политики контроля"),
        ("Users", "Создает, просматривает и сопровождает пользователей и компьютеры", "При изменении состава сотрудников или рабочих мест"),
    ]:
        row = sections_table.add_row().cells
        row[0].text = section
        row[1].text = action
        row[2].text = when
    add_table_font(sections_table)
    add_table_caption(doc, "Таблица 8.3. Типовые ситуации и действия пользователя")
    situations_table = doc.add_table(rows=1, cols=2)
    situations_table.style = "Table Grid"
    situations_table.rows[0].cells[0].text = "Ситуация"
    situations_table.rows[0].cells[1].text = "Действие пользователя"
    for situation, action in [
        ("Не удается войти в систему", "Проверить логин и пароль, затем повторить вход; при сохранении ошибки обратиться к администратору стенда"),
        ("Dashboard показывает нулевые значения", "Убедиться, что на стенд загружены пользователи, компьютеры, агенты и события активности"),
        ("Агент долго не выходит на связь", "Открыть Agents, проверить статус, последнюю активность и при необходимости повторить команду или проверить сеть"),
        ("Нужно разобраться в подозрительном событии", "Открыть Analytics, выбрать период и фильтры, затем перейти к отчету или журналу событий"),
        ("Необходимо изменить правила мониторинга", "Открыть Settings, изменить списки или правила, сохранить настройки и проверить синхронизацию политик"),
        ("Работа завершена", "Выполнить выход из системы, чтобы закрыть административную сессию"),
    ]:
        row = situations_table.add_row().cells
        row[0].text = situation
        row[1].text = action
    add_table_font(situations_table)
    add_centered_picture(
        doc,
        control_plane_img,
        5.6,
        "Рисунок 8.1. Контур взаимодействия администратора, control plane и endpoint-агента",
    )
    add_text_paragraph(
        doc,
        "Представленный сценарий показывает последовательность действий администратора от входа в систему до управления агентами и просмотра аналитики. После проверки учетных данных пользователь работает только с защищенными разделами, а изменения конфигурации и командный контур фиксируются серверными механизмами аудита."
    )

    doc.add_page_break()
    add_heading(doc, "Приложение 9. Скриншоты демонстрационного стенда", level=2)
    add_text_paragraph(
        doc,
        "Скриншоты подготовлены 19.04.2026 с демонстрационного сервера http://2.26.89.86 после проверки авторизации и переходов по основным разделам web-приложения. На момент фиксации стенд находился в начальном состоянии данных: предметные пользователи, агенты, события активности, аномалии и правила оповещений отсутствовали."
    )
    screenshot_rows = [
        (
            "01_login.png",
            "Рисунок 9.1. Страница входа в демонстрационный стенд",
            "Экран подтверждает наличие отдельной страницы авторизации, через которую администратор получает доступ к защищенным разделам системы.",
        ),
        (
            "02_dashboard.png",
            "Рисунок 9.2. Панель Dashboard с текущими показателями стенда",
            "Панель отображает базовые KPI и состояние активности. Нулевые значения соответствуют свежему стенду без загруженной телеметрии.",
        ),
        (
            "03_agents.png",
            "Рисунок 9.3. Раздел управления endpoint-агентами",
            "Раздел Agents предназначен для просмотра подключенных агентов, их политик и командного контура. В текущем состоянии список агентов пуст.",
        ),
        (
            "04_analytics.png",
            "Рисунок 9.4. Раздел аналитики пользовательской активности",
            "Раздел Analytics сохраняет рабочую структуру фильтров и аналитических блоков даже при отсутствии накопленных событий активности.",
        ),
        (
            "05_reports.png",
            "Рисунок 9.5. Раздел отчетности и аналитических выгрузок",
            "Раздел Reports предназначен для формирования сводок и экспорта результатов после появления событий и отчетных проекций.",
        ),
        (
            "06_settings.png",
            "Рисунок 9.6. Раздел системных настроек мониторинга",
            "Раздел Settings отражает доступность параметров безопасности, мониторинга, списков доступа и правил оповещений.",
        ),
        (
            "07_users.png",
            "Рисунок 9.7. Раздел управления пользователями",
            "Раздел Users подтверждает наличие административного маршрута для работы с пользователями и связанными рабочими станциями.",
        ),
    ]
    for filename, caption, description in screenshot_rows:
        image_path = SCREENSHOTS / filename
        if image_path.exists():
            add_centered_picture(doc, image_path, 6.2, caption)
            add_text_paragraph(doc, description)
        else:
            add_text_paragraph(doc, f"Файл скриншота {filename} отсутствует в каталоге материалов и требует повторной фиксации.")

    add_table_caption(doc, "Таблица 9.1. Итоговая фиксация состояния демонстрационного стенда")
    stand_table = doc.add_table(rows=1, cols=2)
    stand_table.style = "Table Grid"
    stand_table.rows[0].cells[0].text = "Проверенный показатель"
    stand_table.rows[0].cells[1].text = "Состояние на 19.04.2026 18:12 МСК"
    for key, value in [
        ("Адрес стенда", "http://2.26.89.86"),
        ("Демонстрационная учетная запись", "admin / admin123"),
        ("Статус сервисов", "healthy для gateway и основных backend-сервисов"),
        ("Пользователи, агенты, активности, аномалии", "0 записей в начальном состоянии стенда"),
        ("Настройки мониторинга", "real-time monitoring и anomaly detection включены"),
    ]:
        row = stand_table.add_row().cells
        row[0].text = key
        row[1].text = value
    add_table_font(stand_table)
    add_text_paragraph(
        doc,
        "Итоговая фиксация показывает, что интерфейсные маршруты доступны, а отсутствие предметных записей является контролируемым начальным состоянием стенда. После наполнения тестовыми пользователями и агентами те же разделы используются для демонстрации динамики активности, аномалий, уведомлений и отчетов."
    )

    format_appendix_content(doc)
    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    build_document()
    print(OUT_DOCX)
