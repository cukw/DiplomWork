# -*- coding: utf-8 -*-
from pathlib import Path
import importlib.util

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/cukw/FinalWork/docs/appendix_materials/Глава_описание_БД_и_таблиц.docx")
SOURCE = Path("/Users/cukw/FinalWork/scripts/build_database_tables_docx.py")

spec = importlib.util.spec_from_file_location("db_tables_source", SOURCE)
source = importlib.util.module_from_spec(spec)
spec.loader.exec_module(source)
DATABASES = source.DATABASES
SUMMARY_PURPOSES = source.SUMMARY_PURPOSES


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=9, bold=False, color=None):
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold, color=color)


def add_paragraph(document, text, style=None, justify=False):
    paragraph = document.add_paragraph(text, style=style)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        style_run(run, size=11)
    return paragraph


def add_table(document, headers, rows, widths, left_cols=(0, 3)):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, size=9, bold=True)

    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = str(value)
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in left_cols else WD_ALIGN_PARAGRAPH.CENTER
                style_paragraph(paragraph, size=8.5)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, width in enumerate(widths):
            row.cells[index].width = width
            set_cell_margins(row.cells[index])

    return table


def setup_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def lower_first(text):
    return text[:1].lower() + text[1:] if text else text


def intro_for_table(table):
    return (
        f"Таблица «{table['name']}» {lower_first(table['purpose'])} "
        "Она фиксирует отдельный объект предметной области и содержит набор полей, "
        "необходимых для его идентификации, описания и дальнейшей обработки в системе. "
        "Состав полей таблицы представлен ниже."
    )


def outro_for_table(table):
    primary_keys = [column[0] for column in table["columns"] if "PK" in column[2]]
    foreign_keys = [column[0] for column in table["columns"] if "FK" in column[2]]
    pk_text = ", ".join(primary_keys) if primary_keys else "служебные идентификаторы"
    if foreign_keys:
        fk_text = f" Поля {', '.join(foreign_keys)} связывают эту таблицу с другими сущностями."
    else:
        fk_text = ""
    return (
        f"Таким образом, структура таблицы «{table['name']}» позволяет хранить данные в упорядоченном виде. "
        f"Ключевым идентификатором выступает {pk_text}.{fk_text} "
        "Остальные поля раскрывают содержательные характеристики записи и используются сервисами приложения "
        "при выполнении операций учета, анализа, уведомления или администрирования."
    )


def build_doc():
    document = Document()
    setup_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Описание базы данных и таблиц информационной системы")
    style_run(run, size=20, bold=True, color="0B2545")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Система учета и анализа компьютерной активности пользователей")
    style_run(run, size=12, italic=True, color="555555")

    document.add_heading("Назначение базы данных в системе", level=1)
    add_paragraph(
        document,
        "База данных в информационной системе предназначена для долговременного и структурированного "
        "хранения сведений, которые возникают при работе пользователей, администраторов и endpoint-агентов. "
        "Без базы данных система могла бы только временно обрабатывать поступающие события, но не смогла бы "
        "сохранять историю активности, формировать отчеты, анализировать отклонения, восстанавливать состояние "
        "агентов и подтверждать действия администратора.",
        justify=True,
    )
    add_paragraph(
        document,
        "В рамках данного проекта база данных выполняет несколько ключевых функций. Во-первых, она хранит "
        "учетные записи, роли и сведения о пользователях. Во-вторых, она фиксирует рабочие станции и события "
        "компьютерной активности. В-третьих, она обеспечивает хранение результатов анализа: аномалий, метрик, "
        "уведомлений и отчетных проекций. В-четвертых, база данных поддерживает управление endpoint-агентами, "
        "их политиками и командами. Поэтому структура хранения напрямую связана с функциональностью всей системы.",
        justify=True,
    )
    add_paragraph(
        document,
        "Так как проект построен на микросервисной архитектуре, данные разделены между несколькими PostgreSQL-"
        "хранилищами. Каждая база данных обслуживает свой сервис и хранит только те таблицы, которые относятся "
        "к его зоне ответственности. Это снижает связанность компонентов, упрощает сопровождение схем и делает "
        "развитие системы более управляемым.",
        justify=True,
    )

    document.add_heading("Краткая сводка по базам данных", level=1)
    summary_rows = [
        (
            db["name"],
            db["physical"],
            SUMMARY_PURPOSES[db["physical"]],
            ", ".join(table["name"] for table in db["tables"]),
        )
        for db in DATABASES
    ]
    add_table(
        document,
        ["База данных", "Физическое имя", "Назначение", "Таблицы"],
        summary_rows,
        [Inches(1.55), Inches(1.0), Inches(2.05), Inches(1.90)],
        left_cols=(0, 2, 3),
    )
    add_paragraph(
        document,
        "Сводная таблица показывает, какие хранилища используются в проекте и какие группы таблиц "
        "входят в их состав. Далее каждая база данных рассматривается отдельно: сначала описывается "
        "ее роль, затем приводится описание входящих в нее таблиц и перечень хранимых полей.",
        justify=True,
    )

    document.add_heading("Описание баз данных и таблиц", level=1)
    for db_index, db in enumerate(DATABASES, start=1):
        if db_index > 1:
            document.add_page_break()
        document.add_heading(f"{db_index}. {db['name']} ({db['physical']})", level=2)
        add_paragraph(document, db["purpose"], justify=True)
        add_paragraph(
            document,
            "В состав этой базы данных входят таблицы, отражающие основные объекты соответствующего "
            "сервиса. Каждая таблица имеет собственное назначение и хранит ограниченный набор полей, "
            "необходимых для выполнения бизнес-логики подсистемы.",
            justify=True,
        )

        for table in db["tables"]:
            document.add_heading(f"Таблица {table['name']}", level=3)
            add_paragraph(document, intro_for_table(table), justify=True)
            add_table(
                document,
                ["Поле", "Тип данных", "Ключ", "Что хранит поле"],
                table["columns"],
                [Inches(1.45), Inches(1.35), Inches(0.95), Inches(2.75)],
                left_cols=(0, 3),
            )
            add_paragraph(document, outro_for_table(table), justify=True)

        add_paragraph(
            document,
            f"В результате {db['name']} обеспечивает хранение данных, необходимых для работы своей "
            "подсистемы, и передает другим компонентам только прикладные результаты через API или "
            "событийные сообщения. Такой подход поддерживает целостность архитектуры и упрощает "
            "дальнейшее сопровождение проекта.",
            justify=True,
        )

    document.add_heading("Вывод", level=1)
    add_paragraph(
        document,
        "База данных в разработанной системе является основой для учета, анализа и последующего "
        "использования сведений о компьютерной активности пользователей. Таблицы распределены по "
        "доменных базам данных в соответствии с ответственностью сервисов. Такое построение позволяет "
        "сохранять историю активности, контролировать пользователей и устройства, управлять агентами, "
        "выявлять аномалии, формировать уведомления и строить отчетность. Наличие описанных полей в "
        "каждой таблице обеспечивает полноту хранения и возможность дальнейшей обработки данных в рамках "
        "информационной системы.",
        justify=True,
    )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Описание БД и таблиц")
    style_run(run, size=9, color="555555")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
