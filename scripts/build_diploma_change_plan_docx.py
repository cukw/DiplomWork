from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "diploma_change_plan_ru.md"
OUTPUT = ROOT / "docs" / "Правки_к_диплому_по_новым_добавлениям.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str = "D9E2F3") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "2")
        node.set(qn("w:color"), color)
        borders.append(node)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Heading 1", 17, "1F4E79", 14, 7),
        ("Heading 2", 13, "2F5597", 12, 5),
        ("Heading 3", 11.5, "1F4E79", 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Правки к диплому с учетом новых добавлений")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4E79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Отдельный файл: что добавить, что удалить и после какого места вставлять")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("666666")


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string("404040")
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown: str) -> Document:
    doc = Document()
    configure_document(doc)
    add_title(doc)

    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.05)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph_shading(paragraph, "F3F6FA")
        set_paragraph_border(paragraph)
        run = paragraph.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        code_lines = []

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, number_match.group(2))
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:])
            continue

        paragraph = doc.add_paragraph()
        add_inline_markdown(paragraph, stripped)

    if in_code:
        flush_code()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Документ подготовлен как список точечных правок к дипломной работе")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("808080")

    return doc


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = markdown_to_docx(markdown)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
