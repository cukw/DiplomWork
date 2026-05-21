# -*- coding: utf-8 -*-
from pathlib import Path
import importlib.util

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/cukw/FinalWork")
SOURCE = ROOT / "scripts" / "build_database_tables_docx.py"
OUT_DOCX = ROOT / "docs" / "Глава_2_1_Проектирование_физической_модели_БД.docx"
OUT_MD = ROOT / "docs" / "Глава_2_1_Проектирование_физической_модели_БД.md"


spec = importlib.util.spec_from_file_location("db_tables_source", SOURCE)
source = importlib.util.module_from_spec(spec)
spec.loader.exec_module(source)
DATABASES = source.DATABASES
SUMMARY_PURPOSES = source.SUMMARY_PURPOSES


TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
FIELD_WIDTHS = [0.33, 1.85, 1.40, 1.15, 0.72, 1.05]
SUMMARY_WIDTHS = [1.65, 1.05, 2.05, 1.75]


def lower_first(text):
    return text[:1].lower() + text[1:] if text else text


def clean_description(text):
    return (text or "").strip().rstrip(".")


def is_required(raw_constraint):
    constraint = raw_constraint.upper()
    return "PK" in constraint or "NOT NULL" in constraint


def constraint_text(raw_constraint):
    constraint = raw_constraint.upper()
    parts = []
    if "PK" in constraint:
        parts.append("первичный ключ")
    if "FK" in constraint:
        parts.append("внешний ключ")
    if "UNIQUE PART" in constraint:
        parts.append("часть уникального индекса")
    elif "UNIQUE" in constraint:
        parts.append("уникальное значение")
    return "; ".join(parts) if parts else "-"


def field_rows(table):
    rows = []
    for index, (identifier, data_type, constraint, description) in enumerate(table["columns"], start=1):
        rows.append(
            (
                index,
                clean_description(description),
                identifier,
                data_type,
                "Да" if is_required(constraint) else "Нет",
                constraint_text(constraint),
            )
        )
    return rows


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, width_dxa, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, **CELL_MARGINS_DXA)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def style_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, size=11, bold=False, italic=False, color=None):
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold, italic=italic, color=color)


def add_paragraph(document, text, style=None, justify=True):
    paragraph = document.add_paragraph(text, style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        style_run(run, size=11)
    return paragraph


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    style_run(run, size=10, italic=True, color="333333")
    return paragraph


def add_table(document, headers, rows, widths_in, left_cols=()):
    widths = [int(round(width * 1440)) for width in widths_in]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(paragraph, size=8.5, bold=True)

    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in left_cols else WD_ALIGN_PARAGRAPH.CENTER
                style_paragraph(paragraph, size=8)

    set_table_geometry(table, TABLE_WIDTH_DXA, widths)
    return table


def setup_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_markdown():
    lines = [
        "2.1 Проектирование физической модели базы данных",
        "",
        "Физическая модель базы данных разработана с учетом микросервисной архитектуры информационной системы мониторинга активности пользователей. В проекте не используется единое монолитное хранилище: данные разделены между сервисами AuthService, UserService, ActivityService, AgentManagementService, MetricsService, NotificationService, ReportService и runtime-хранилищем API Gateway. Такое решение снижает связанность компонентов, упрощает сопровождение схем и позволяет изолировать транзакционные нагрузки разных подсистем.",
        "",
        "Основными объектами хранения являются учетные записи и роли пользователей, бизнес-профили сотрудников, рабочие станции, события активности, выявленные аномалии, политики endpoint-агентов, команды управления, метрики, уведомления, отчеты и журнал административных действий. Ниже приведена структура баз данных и таблиц в формате словаря данных: для каждого поля указаны его наименование, технический идентификатор, тип данных, обязательность заполнения и ключевое ограничение.",
        "",
        "Таблица 2.1.1 - Состав баз данных проекта",
        "",
        "| № | База данных | Физическое имя | Назначение | Основные таблицы |",
        "|---|---|---|---|---|",
    ]
    for index, db in enumerate(DATABASES, start=1):
        tables = ", ".join(table["name"] for table in db["tables"])
        lines.append(
            f"| {index} | {db['name']} | `{db['physical']}` | {SUMMARY_PURPOSES[db['physical']]} | {tables} |"
        )

    table_number = 2
    for db in DATABASES:
        lines.extend(["", f"#### {db['name']} ({db['physical']})", "", db["purpose"], ""])
        lines.append(
            "В данной базе данных выделены таблицы, которые соответствуют зоне ответственности сервиса. "
            "Их структура представлена ниже."
        )
        for table in db["tables"]:
            lines.extend(
                [
                    "",
                    f"Таблица 2.1.{table_number} - Структура таблицы `{table['name']}`",
                    "",
                    "| № | Название | Идентификатор | Тип | Не пусто | Ограничение |",
                    "|---|---|---|---|---|---|",
                ]
            )
            for row in field_rows(table):
                lines.append(
                    f"| {row[0]} | {row[1]} | `{row[2]}` | `{row[3]}` | {row[4]} | {row[5]} |"
                )
            lines.extend(
                [
                    "",
                    f"Таблица `{table['name']}` {lower_first(table['purpose'])} "
                    "Набор полей обеспечивает идентификацию записи, хранение основных характеристик объекта "
                    "и дальнейшую обработку данных сервисами приложения.",
                ]
            )
            table_number += 1

    lines.extend(
        [
            "",
            "Таким образом, физическая модель базы данных покрывает полный цикл работы системы: регистрацию и авторизацию пользователей, учет рабочих станций, прием событий активности, выявление аномалий, построение агрегатов, доставку уведомлений, формирование отчетов и управление endpoint-агентами. Служебные таблицы outbox, inbox и DLQ повышают надежность событийной обработки, а разделение хранилищ по сервисам соответствует выбранной микросервисной архитектуре.",
            "",
        ]
    )
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_docx():
    document = Document()
    setup_document(document)

    document.add_heading("2.1 Проектирование физической модели базы данных", level=1)
    add_paragraph(
        document,
        "Физическая модель базы данных разработана с учетом микросервисной архитектуры информационной "
        "системы мониторинга активности пользователей. В проекте не используется единое монолитное "
        "хранилище: данные разделены между сервисами AuthService, UserService, ActivityService, "
        "AgentManagementService, MetricsService, NotificationService, ReportService и runtime-хранилищем "
        "API Gateway. Такое решение снижает связанность компонентов, упрощает сопровождение схем и "
        "позволяет изолировать транзакционные нагрузки разных подсистем.",
    )
    add_paragraph(
        document,
        "Основными объектами хранения являются учетные записи и роли пользователей, бизнес-профили "
        "сотрудников, рабочие станции, события активности, выявленные аномалии, политики endpoint-агентов, "
        "команды управления, метрики, уведомления, отчеты и журнал административных действий. Ниже приведена "
        "структура баз данных и таблиц в формате словаря данных: для каждого поля указаны его наименование, "
        "технический идентификатор, тип данных, обязательность заполнения и ключевое ограничение.",
    )

    add_caption(document, "Таблица 2.1.1 - Состав баз данных проекта")
    summary_rows = [
        (
            index,
            db["name"],
            db["physical"],
            SUMMARY_PURPOSES[db["physical"]],
        )
        for index, db in enumerate(DATABASES, start=1)
    ]
    add_table(
        document,
        ["№", "База данных", "Физическое имя", "Назначение"],
        summary_rows,
        [0.35, 1.85, 1.15, 3.15],
        left_cols=(1, 3),
    )

    table_number = 2
    for db_index, db in enumerate(DATABASES, start=1):
        if db_index > 1:
            document.add_page_break()
        document.add_heading(f"{db['name']} ({db['physical']})", level=2)
        add_paragraph(document, db["purpose"])
        add_paragraph(
            document,
            "В данной базе данных выделены таблицы, которые соответствуют зоне ответственности сервиса. "
            "Их структура представлена ниже.",
        )
        for table in db["tables"]:
            document.add_heading(f"Таблица {table['name']}", level=3)
            add_paragraph(
                document,
                f"Таблица «{table['name']}» {lower_first(table['purpose'])} "
                "Состав полей приведен в словаре данных.",
            )
            add_caption(document, f"Таблица 2.1.{table_number} - Структура таблицы {table['name']}")
            add_table(
                document,
                ["№", "Название", "Идентификатор", "Тип", "Не пусто", "Ограничение"],
                field_rows(table),
                FIELD_WIDTHS,
                left_cols=(1, 2, 5),
            )
            add_paragraph(
                document,
                f"Набор полей таблицы «{table['name']}» обеспечивает идентификацию записи, хранение "
                "основных характеристик объекта и дальнейшую обработку данных сервисами приложения.",
            )
            table_number += 1

    document.add_heading("Вывод по пункту 2.1", level=2)
    add_paragraph(
        document,
        "Таким образом, физическая модель базы данных покрывает полный цикл работы системы: регистрацию "
        "и авторизацию пользователей, учет рабочих станций, прием событий активности, выявление аномалий, "
        "построение агрегатов, доставку уведомлений, формирование отчетов и управление endpoint-агентами. "
        "Служебные таблицы outbox, inbox и DLQ повышают надежность событийной обработки, а разделение "
        "хранилищ по сервисам соответствует выбранной микросервисной архитектуре.",
    )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Глава 2.1. Проектирование физической модели базы данных")
    style_run(run, size=9, color="555555")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)


def main():
    build_markdown()
    build_docx()
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
