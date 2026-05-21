# -*- coding: utf-8 -*-
from pathlib import Path
import importlib.util

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/cukw/FinalWork/docs/appendix_materials/Краткая_сводка_БД_и_приложения.docx")
SOURCE = Path("/Users/cukw/FinalWork/scripts/build_database_tables_docx.py")


spec = importlib.util.spec_from_file_location("db_tables_source", SOURCE)
source = importlib.util.module_from_spec(spec)
spec.loader.exec_module(source)
DATABASES = source.DATABASES
SUMMARY_PURPOSES = source.SUMMARY_PURPOSES


APPENDIX_LETTERS = ["А", "Б", "В", "Г", "Д", "Е", "Ж", "З"]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=10, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def apply_paragraph_style(paragraph, size=9, bold=False, color=None):
    for run in paragraph.runs:
        style_run(run, size=size, bold=bold, color=color)


def add_paragraph(document, text, style=None):
    paragraph = document.add_paragraph(text, style=style)
    for run in paragraph.runs:
        style_run(run, size=11)
    return paragraph


def add_table(document, headers, rows, widths, left_cols=(1, 2, 3)):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_style(paragraph, size=9, bold=True)

    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            cell.text = str(value)
            cell.width = widths[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index in left_cols else WD_ALIGN_PARAGRAPH.CENTER
                apply_paragraph_style(paragraph, size=8.5)

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for index, width in enumerate(widths):
            row.cells[index].width = width
            set_cell_margins(row.cells[index])
    return table


def setup_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def tuple_for_table(table):
    fields = ", ".join(column[0] for column in table["columns"])
    return f"{table['name']}({fields})"


def main():
    document = Document()
    setup_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Краткая сводка по базам данных и оформлению приложений")
    style_run(run, size=20, bold=True, color="0B2545")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Система учета и анализа компьютерной активности пользователей")
    style_run(run, size=12, italic=True, color="555555")

    add_paragraph(
        document,
        "В информационной системе используется несколько PostgreSQL-хранилищ, каждое из которых "
        "закреплено за отдельным сервисом. Такой подход позволяет разделить ответственность между "
        "подсистемами: аутентификацией, учетом пользователей, сбором активности, управлением агентами, "
        "метриками, уведомлениями, отчетностью и runtime-настройками API Gateway.",
    )
    add_paragraph(
        document,
        "Ниже приведена краткая сводка по каждой базе данных. Для каждой БД указано назначение, "
        "таблицы в виде таблицы, кортежи отношений и готовая формулировка, которую можно вставить "
        "в основной текст диплома при ссылке на приложение.",
    )

    document.add_heading("Общая сводка по базам данных", level=1)
    summary_rows = []
    for db, letter in zip(DATABASES, APPENDIX_LETTERS):
        summary_rows.append(
            (
                f"Приложение {letter}",
                db["name"],
                db["physical"],
                SUMMARY_PURPOSES[db["physical"]],
            )
        )
    add_table(
        document,
        ["Приложение", "База данных", "Физическое имя", "Краткое назначение"],
        summary_rows,
        [Inches(1.05), Inches(1.65), Inches(1.15), Inches(2.65)],
        left_cols=(1, 3),
    )
    add_paragraph(
        document,
        "Эту таблицу можно использовать как вводную сводку перед подробным описанием приложений. "
        "Далее каждая база данных раскрывается отдельно, чтобы было понятно, какой материал должен "
        "быть вынесен в соответствующее приложение.",
    )

    for index, (db, letter) in enumerate(zip(DATABASES, APPENDIX_LETTERS), start=1):
        if index > 1:
            document.add_page_break()
        document.add_heading(f"{index}. {db['name']} ({db['physical']})", level=1)
        add_paragraph(document, db["purpose"])

        document.add_heading("Таблицы базы данных", level=2)
        table_rows = []
        for number, table in enumerate(db["tables"], start=1):
            key_columns = ", ".join(column[0] for column in table["columns"] if "PK" in column[2])
            table_rows.append((number, table["name"], key_columns, table["purpose"]))
        add_table(
            document,
            ["№", "Таблица", "Первичный ключ", "Назначение таблицы"],
            table_rows,
            [Inches(0.45), Inches(1.45), Inches(1.30), Inches(3.30)],
            left_cols=(1, 2, 3),
        )
        add_paragraph(
            document,
            "Представленный перечень показывает, какие отношения входят в состав базы данных и какую "
            "роль они выполняют в обработке данных системы. Для раскрытия структуры этих отношений ниже "
            "приведены их кортежи.",
        )

        document.add_heading("Кортежи отношений", level=2)
        tuple_rows = [(table["name"], tuple_for_table(table)) for table in db["tables"]]
        add_table(
            document,
            ["Отношение", "Кортеж отношения"],
            tuple_rows,
            [Inches(1.45), Inches(5.05)],
            left_cols=(0, 1),
        )
        add_paragraph(
            document,
            "Кортежи отношений показывают состав атрибутов каждой таблицы и могут быть использованы "
            "как краткое приложение к разделу проектирования базы данных.",
        )

        document.add_heading("Как вставить приложение в текст работы", level=2)
        add_paragraph(
            document,
            f"Рекомендуемый заголовок приложения: «Приложение {letter}. Структура таблиц базы данных "
            f"{db['name']}». В основной части работы после описания соответствующего сервиса можно "
            "вставить следующую формулировку:",
        )
        quote = document.add_paragraph()
        quote.paragraph_format.left_indent = Inches(0.3)
        quote.paragraph_format.right_indent = Inches(0.2)
        quote.paragraph_format.space_before = Pt(4)
        quote.paragraph_format.space_after = Pt(8)
        run = quote.add_run(
            f"Структура таблиц базы данных {db['name']} представлена в приложении {letter}. "
            f"В приложении приведены основные отношения, их кортежи и назначение таблиц, "
            "используемых данной подсистемой."
        )
        style_run(run, size=10, italic=True, color="333333")
        add_paragraph(
            document,
            "Такая ссылка позволяет не перегружать основной текст подробным перечислением полей, "
            "но при этом сохраняет связь между проектным описанием сервиса и фактической структурой "
            "его базы данных.",
        )

    document.add_page_break()
    document.add_heading("Итоговая рекомендация по оформлению", level=1)
    add_paragraph(
        document,
        "В основной главе дипломной работы целесообразно оставить краткое описание назначения баз "
        "данных и указать, что детальная структура вынесена в приложения. В приложениях следует "
        "разместить таблицы с отношениями и кортежами, а также при необходимости добавить полное "
        "описание столбцов. Такой подход делает основной текст более плавным, а технические детали "
        "остаются доступными для проверки.",
    )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Краткая сводка БД и приложения")
    style_run(run, size=9, color="555555")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
