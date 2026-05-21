# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/cukw/FinalWork/docs/appendix_materials/Описание_таблиц_базы_данных.docx")


DATABASES = [
    {
        "name": "База данных AuthService",
        "physical": "auth",
        "purpose": (
            "База данных сервиса аутентификации предназначена для хранения учетных записей, "
            "ролей пользователей и активных сессий. Она обеспечивает вход в систему, проверку "
            "прав пользователя и поддержку жизненного цикла токенов доступа."
        ),
        "tables": [
            {
                "name": "roles",
                "purpose": "Хранит роли пользователей системы и их описание.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код роли."),
                    ("name", "VARCHAR(50)", "UNIQUE, NOT NULL", "Название роли: admin, user, moderator, auditor."),
                    ("description", "TEXT", "", "Текстовое описание роли."),
                    ("created_at", "TIMESTAMP", "", "Дата создания записи."),
                ],
            },
            {
                "name": "auth_users",
                "purpose": "Хранит учетные записи, используемые для входа в административную панель.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код учетной записи."),
                    ("username", "VARCHAR(100)", "UNIQUE, NOT NULL", "Логин пользователя."),
                    ("password_hash", "VARCHAR(255)", "NOT NULL", "Хеш пароля."),
                    ("email", "VARCHAR(255)", "UNIQUE", "Электронная почта пользователя."),
                    ("role_id", "INTEGER", "FK", "Ссылка на роль пользователя."),
                    ("last_login", "TIMESTAMP", "", "Дата и время последнего входа."),
                    ("is_active", "BOOLEAN", "", "Признак активности учетной записи."),
                    ("created_at", "TIMESTAMP", "", "Дата создания учетной записи."),
                ],
            },
            {
                "name": "sessions",
                "purpose": "Хранит пользовательские сессии и refresh-токены.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код сессии."),
                    ("user_id", "INTEGER", "FK, NOT NULL", "Ссылка на учетную запись."),
                    ("token_hash", "VARCHAR(255)", "UNIQUE, NOT NULL", "Хеш токена сессии."),
                    ("expires_at", "TIMESTAMP", "NOT NULL", "Срок действия сессии."),
                    ("created_at", "TIMESTAMP", "", "Дата создания сессии."),
                ],
            },
        ],
    },
    {
        "name": "База данных UserService",
        "physical": "users",
        "purpose": (
            "База данных сервиса пользователей хранит бизнес-профили пользователей и сведения "
            "об их рабочих станциях. Она используется для привязки событий активности к конкретному "
            "сотруднику и компьютеру."
        ),
        "tables": [
            {
                "name": "users",
                "purpose": "Хранит бизнес-профили пользователей системы.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код пользователя."),
                    ("auth_user_id", "INTEGER", "UNIQUE", "Код учетной записи в сервисе аутентификации."),
                    ("full_name", "VARCHAR(255)", "", "ФИО пользователя."),
                    ("department", "VARCHAR(100)", "", "Подразделение пользователя."),
                    ("created_at", "TIMESTAMP", "", "Дата создания записи."),
                ],
            },
            {
                "name": "computers",
                "purpose": "Хранит рабочие станции пользователей и их техническое состояние.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код компьютера."),
                    ("user_id", "INTEGER", "FK, UNIQUE, NOT NULL", "Ссылка на пользователя; связь один к одному."),
                    ("hostname", "VARCHAR(255)", "NOT NULL", "Имя рабочей станции."),
                    ("os_version", "VARCHAR(100)", "", "Версия операционной системы."),
                    ("ip_address", "INET", "", "IP-адрес компьютера."),
                    ("mac_address", "VARCHAR(17)", "UNIQUE", "MAC-адрес сетевого интерфейса."),
                    ("status", "VARCHAR(20)", "", "Статус компьютера: active, disabled, retired."),
                    ("last_seen", "TIMESTAMP", "", "Время последней активности."),
                    ("created_at", "TIMESTAMP", "", "Дата создания записи."),
                ],
            },
        ],
    },
    {
        "name": "База данных ActivityService",
        "physical": "activities",
        "purpose": (
            "База данных сервиса активности является центральным хранилищем событий компьютерной "
            "активности. Она сохраняет поступающие от endpoint-агентов события, результаты анализа "
            "риска, выявленные аномалии, архивные записи и outbox-события для публикации в RabbitMQ."
        ),
        "tables": [
            {
                "name": "activities",
                "purpose": "Хранит события компьютерной активности пользователей.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код события активности."),
                    ("computer_id", "INTEGER", "NOT NULL", "Код компьютера, на котором произошло событие."),
                    ("timestamp", "TIMESTAMPTZ", "", "Дата и время события."),
                    ("activity_type", "VARCHAR(50)", "NOT NULL", "Тип активности: процесс, сайт, файл и т.д."),
                    ("details", "JSONB", "", "Дополнительные сведения о событии."),
                    ("duration_ms", "INTEGER", "", "Длительность действия в миллисекундах."),
                    ("url", "VARCHAR(500)", "", "URL-адрес при браузерной активности."),
                    ("process_name", "VARCHAR(255)", "", "Имя процесса."),
                    ("is_blocked", "BOOLEAN", "", "Признак блокировки действия."),
                    ("risk_score", "NUMERIC(5,2)", "", "Оценка риска от 0 до 100."),
                    ("synced", "BOOLEAN", "", "Признак синхронизации события."),
                    ("user_id", "BIGINT", "", "Код пользователя."),
                    ("agent_id", "BIGINT", "", "Код endpoint-агента."),
                    ("agent_version", "VARCHAR(50)", "", "Версия агента."),
                    ("device_name", "VARCHAR(255)", "", "Имя устройства."),
                    ("collector", "VARCHAR(100)", "", "Коллектор, сформировавший событие."),
                    ("event_id", "VARCHAR(100)", "", "Внешний идентификатор события."),
                    ("sequence", "BIGINT", "", "Порядковый номер события от агента."),
                    ("batch_id", "VARCHAR(100)", "", "Идентификатор пакета синхронизации."),
                    ("source_platform", "VARCHAR(50)", "", "Платформа источника события."),
                ],
            },
            {
                "name": "activities_archive",
                "purpose": "Хранит архивные копии событий активности после применения политики хранения.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код архивной записи."),
                    ("original_activity_id", "BIGINT", "UNIQUE, NOT NULL", "Код исходного события активности."),
                    ("computer_id", "INTEGER", "NOT NULL", "Код компьютера."),
                    ("timestamp", "TIMESTAMPTZ", "NOT NULL", "Дата и время события."),
                    ("activity_type", "VARCHAR(50)", "NOT NULL", "Тип активности."),
                    ("details", "JSONB", "", "Дополнительные сведения."),
                    ("duration_ms", "INTEGER", "", "Длительность действия."),
                    ("url", "VARCHAR(500)", "", "URL-адрес."),
                    ("process_name", "VARCHAR(255)", "", "Имя процесса."),
                    ("is_blocked", "BOOLEAN", "", "Признак блокировки."),
                    ("risk_score", "NUMERIC(5,2)", "", "Оценка риска."),
                    ("synced", "BOOLEAN", "", "Признак синхронизации."),
                    ("user_id", "BIGINT", "", "Код пользователя."),
                    ("agent_id", "BIGINT", "", "Код агента."),
                    ("agent_version", "VARCHAR(50)", "", "Версия агента."),
                    ("device_name", "VARCHAR(255)", "", "Имя устройства."),
                    ("collector", "VARCHAR(100)", "", "Название коллектора."),
                    ("event_id", "VARCHAR(100)", "", "Идентификатор события."),
                    ("sequence", "BIGINT", "", "Порядковый номер события."),
                    ("batch_id", "VARCHAR(100)", "", "Идентификатор пакета."),
                    ("source_platform", "VARCHAR(50)", "", "Платформа источника."),
                    ("archived_at", "TIMESTAMPTZ", "", "Дата помещения записи в архив."),
                ],
            },
            {
                "name": "anomalies",
                "purpose": "Хранит аномалии, выявленные при анализе событий активности.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код аномалии."),
                    ("activity_id", "BIGINT", "FK, NOT NULL", "Ссылка на событие активности."),
                    ("type", "VARCHAR(100)", "NOT NULL", "Тип аномалии."),
                    ("description", "TEXT", "", "Описание выявленного отклонения."),
                    ("detected_at", "TIMESTAMPTZ", "", "Дата обнаружения аномалии."),
                ],
            },
            {
                "name": "activity_outbox",
                "purpose": "Фиксирует доменные события, которые должны быть опубликованы в RabbitMQ.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код outbox-сообщения."),
                    ("event_type", "VARCHAR(128)", "NOT NULL", "Тип публикуемого события."),
                    ("activity_id", "BIGINT", "", "Код связанного события активности."),
                    ("payload", "JSONB", "NOT NULL", "Полезная нагрузка сообщения."),
                    ("headers", "JSONB", "", "Заголовки сообщения."),
                    ("attempt_count", "INTEGER", "NOT NULL", "Количество попыток публикации."),
                    ("available_at", "TIMESTAMPTZ", "NOT NULL", "Дата доступности сообщения для отправки."),
                    ("processed_at", "TIMESTAMPTZ", "", "Дата успешной обработки."),
                    ("last_error", "TEXT", "", "Последняя ошибка публикации."),
                    ("created_at", "TIMESTAMPTZ", "NOT NULL", "Дата создания сообщения."),
                ],
            },
        ],
    },
    {
        "name": "База данных AgentManagementService",
        "physical": "agents",
        "purpose": (
            "База данных сервиса управления агентами хранит состояние endpoint-агентов, их политики, "
            "команды управления и сведения о синхронизации. Она обеспечивает централизованную контрольную "
            "плоскость для рабочих станций пользователей."
        ),
        "tables": [
            {
                "name": "agents",
                "purpose": "Хранит зарегистрированные endpoint-агенты и их текущее состояние.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код агента."),
                    ("computer_id", "INTEGER", "UNIQUE", "Код компьютера, на котором установлен агент."),
                    ("version", "VARCHAR(20)", "NOT NULL", "Версия агента."),
                    ("status", "VARCHAR(20)", "NOT NULL", "Статус агента: online, offline, updating."),
                    ("last_heartbeat", "TIMESTAMP", "", "Время последнего heartbeat."),
                    ("config_version", "VARCHAR(20)", "", "Версия конфигурации."),
                    ("offline_since", "TIMESTAMP", "", "Дата перехода агента в offline."),
                    ("desired_version", "VARCHAR(20)", "", "Требуемая версия агента."),
                    ("desired_version_set_at", "TIMESTAMP", "", "Дата назначения требуемой версии."),
                    ("health_json", "TEXT", "NOT NULL", "Диагностическое состояние агента."),
                    ("queue_size", "INTEGER", "NOT NULL", "Размер локальной очереди агента."),
                    ("last_collected_at", "TIMESTAMP", "", "Дата последнего сбора данных."),
                    ("last_sent_at", "TIMESTAMP", "", "Дата последней отправки данных."),
                    ("last_error", "VARCHAR(500)", "NOT NULL", "Последняя ошибка агента."),
                    ("policy_version", "VARCHAR(50)", "", "Текущая версия политики."),
                    ("capabilities_json", "TEXT", "NOT NULL", "Возможности агента."),
                    ("collector_statuses_json", "TEXT", "NOT NULL", "Статусы коллекторов."),
                    ("source_platform", "VARCHAR(50)", "", "Платформа агента."),
                ],
            },
            {
                "name": "sync_batches",
                "purpose": "Фиксирует пакеты синхронизации событий, отправляемые агентами.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код пакета синхронизации."),
                    ("agent_id", "INTEGER", "FK, NOT NULL", "Ссылка на агента."),
                    ("batch_id", "VARCHAR(100)", "NOT NULL", "Внешний идентификатор пакета."),
                    ("status", "VARCHAR(20)", "NOT NULL", "Статус пакета: pending, success, failed."),
                    ("synced_at", "TIMESTAMP", "", "Дата синхронизации."),
                    ("records_count", "INTEGER", "", "Количество записей в пакете."),
                ],
            },
            {
                "name": "agent_policies",
                "purpose": "Хранит актуальные политики мониторинга для endpoint-агентов.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код политики."),
                    ("agent_id", "INTEGER", "FK, UNIQUE, NOT NULL", "Ссылка на агента."),
                    ("computer_id", "INTEGER", "NOT NULL", "Код компьютера."),
                    ("policy_version", "VARCHAR(50)", "NOT NULL", "Версия политики."),
                    ("collection_interval_sec", "INTEGER", "NOT NULL", "Интервал сбора данных."),
                    ("heartbeat_interval_sec", "INTEGER", "NOT NULL", "Интервал heartbeat."),
                    ("flush_interval_sec", "INTEGER", "NOT NULL", "Интервал отправки накопленных данных."),
                    ("enable_process_collection", "BOOLEAN", "NOT NULL", "Включение сбора процессов."),
                    ("enable_browser_collection", "BOOLEAN", "NOT NULL", "Включение сбора браузерной активности."),
                    ("enable_active_window_collection", "BOOLEAN", "NOT NULL", "Включение сбора активного окна."),
                    ("enable_idle_collection", "BOOLEAN", "NOT NULL", "Включение контроля простоя."),
                    ("idle_threshold_sec", "INTEGER", "NOT NULL", "Порог простоя пользователя."),
                    ("browser_poll_interval_sec", "INTEGER", "NOT NULL", "Интервал опроса браузера."),
                    ("process_snapshot_limit", "INTEGER", "NOT NULL", "Лимит снимка процессов."),
                    ("high_risk_threshold", "REAL", "NOT NULL", "Порог высокого риска."),
                    ("auto_lock_enabled", "BOOLEAN", "NOT NULL", "Признак автоматической блокировки."),
                    ("admin_blocked", "BOOLEAN", "NOT NULL", "Признак блокировки администратором."),
                    ("blocked_reason", "VARCHAR(500)", "", "Причина блокировки."),
                    ("browsers_json", "TEXT", "NOT NULL", "Список контролируемых браузеров."),
                    ("whitelist_json", "TEXT", "", "Список разрешенных шаблонов."),
                    ("blacklist_json", "TEXT", "", "Список запрещенных шаблонов."),
                    ("updated_at", "TIMESTAMP", "NOT NULL", "Дата обновления политики."),
                ],
            },
            {
                "name": "agent_policy_versions",
                "purpose": "Хранит историю изменений политик мониторинга.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код версии политики."),
                    ("agent_id", "INTEGER", "FK, NOT NULL", "Ссылка на агента."),
                    ("policy_version", "VARCHAR(50)", "NOT NULL", "Номер версии политики."),
                    ("change_type", "VARCHAR(20)", "NOT NULL", "Тип изменения: create, update, delete, rollback."),
                    ("changed_by", "VARCHAR(100)", "NOT NULL", "Инициатор изменения."),
                    ("snapshot_json", "TEXT", "NOT NULL", "Снимок политики."),
                    ("created_at", "TIMESTAMP", "NOT NULL", "Дата создания версии."),
                ],
            },
            {
                "name": "agent_commands",
                "purpose": "Хранит команды, отправляемые endpoint-агентам.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код команды."),
                    ("agent_id", "INTEGER", "FK, NOT NULL", "Ссылка на агента."),
                    ("command_key", "VARCHAR(100)", "NOT NULL", "Уникальный ключ команды для агента."),
                    ("type", "VARCHAR(50)", "NOT NULL", "Тип команды."),
                    ("payload_json", "TEXT", "NOT NULL", "Параметры команды."),
                    ("status", "VARCHAR(20)", "NOT NULL", "Статус выполнения команды."),
                    ("requested_by", "VARCHAR(100)", "NOT NULL", "Инициатор команды."),
                    ("result_message", "VARCHAR(500)", "NOT NULL", "Результат выполнения."),
                    ("delivery_attempts", "INTEGER", "NOT NULL", "Количество попыток доставки."),
                    ("max_delivery_attempts", "INTEGER", "NOT NULL", "Максимальное число попыток."),
                    ("last_dispatch_at", "TIMESTAMP", "", "Дата последней отправки."),
                    ("next_retry_at", "TIMESTAMP", "", "Дата следующей попытки."),
                    ("timeout_at", "TIMESTAMP", "", "Время истечения команды."),
                    ("dead_letter_reason", "VARCHAR(500)", "NOT NULL", "Причина перевода в DLQ."),
                    ("created_at", "TIMESTAMP", "NOT NULL", "Дата создания команды."),
                    ("acknowledged_at", "TIMESTAMP", "", "Дата подтверждения команды агентом."),
                ],
            },
            {
                "name": "agent_command_dlq",
                "purpose": "Хранит команды, которые не удалось доставить или выполнить.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код записи DLQ."),
                    ("agent_command_id", "INTEGER", "FK, UNIQUE, NOT NULL", "Ссылка на исходную команду."),
                    ("agent_id", "INTEGER", "NOT NULL", "Код агента."),
                    ("command_key", "VARCHAR(100)", "NOT NULL", "Ключ команды."),
                    ("type", "VARCHAR(50)", "NOT NULL", "Тип команды."),
                    ("payload_json", "TEXT", "NOT NULL", "Параметры команды."),
                    ("reason", "VARCHAR(500)", "NOT NULL", "Причина ошибки."),
                    ("delivery_attempts", "INTEGER", "NOT NULL", "Количество попыток доставки."),
                    ("failed_at", "TIMESTAMP", "NOT NULL", "Дата фиксации ошибки."),
                ],
            },
        ],
    },
    {
        "name": "База данных MetricsService",
        "physical": "metrics",
        "purpose": (
            "База данных сервиса метрик хранит настраиваемые правила контроля, whitelist/blacklist "
            "и агрегированные показатели активности. Она используется для аналитики и быстрого получения "
            "сводных метрик без повторного пересчета всех исходных событий."
        ),
        "tables": [
            {
                "name": "metrics",
                "purpose": "Хранит настраиваемые метрики и правила контроля.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код метрики."),
                    ("user_id", "INTEGER", "", "Код пользователя, если правило персональное."),
                    ("type", "VARCHAR(50)", "NOT NULL", "Тип метрики: process, site, file, generic."),
                    ("config", "JSONB", "NOT NULL", "Конфигурация метрики."),
                    ("is_active", "BOOLEAN", "", "Признак активности метрики."),
                    ("updated_at", "TIMESTAMP", "", "Дата обновления."),
                ],
            },
            {
                "name": "whitelists",
                "purpose": "Хранит разрешенные шаблоны для правил контроля.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код записи whitelist."),
                    ("metric_id", "INTEGER", "FK, NOT NULL", "Ссылка на метрику."),
                    ("pattern", "VARCHAR(500)", "NOT NULL", "Разрешенный шаблон."),
                    ("action", "VARCHAR(20)", "", "Действие при совпадении, обычно allow."),
                    ("created_at", "TIMESTAMP", "", "Дата создания записи."),
                ],
            },
            {
                "name": "blacklists",
                "purpose": "Хранит запрещенные шаблоны для правил контроля.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код записи blacklist."),
                    ("metric_id", "INTEGER", "FK, NOT NULL", "Ссылка на метрику."),
                    ("pattern", "VARCHAR(500)", "NOT NULL", "Запрещенный шаблон."),
                    ("action", "VARCHAR(20)", "", "Действие при совпадении, обычно block."),
                    ("created_at", "TIMESTAMP", "", "Дата создания записи."),
                ],
            },
            {
                "name": "processed_event_inbox",
                "purpose": "Фиксирует обработанные сообщения RabbitMQ для защиты от дублей.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код inbox-записи."),
                    ("consumer", "VARCHAR(128)", "NOT NULL", "Название потребителя."),
                    ("event_key", "VARCHAR(256)", "NOT NULL", "Ключ события."),
                    ("message_id", "VARCHAR(128)", "", "Идентификатор сообщения."),
                    ("processed_at", "TIMESTAMPTZ", "NOT NULL", "Дата обработки события."),
                ],
            },
            {
                "name": "activity_event_rollups",
                "purpose": "Хранит агрегаты событий активности по дате, компьютеру и типу активности.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код агрегата."),
                    ("bucket_date", "DATE", "UNIQUE PART", "Дата агрегирования."),
                    ("computer_id", "INTEGER", "UNIQUE PART, NOT NULL", "Код компьютера."),
                    ("activity_type", "VARCHAR(100)", "UNIQUE PART, NOT NULL", "Тип активности."),
                    ("total_count", "BIGINT", "NOT NULL", "Общее количество событий."),
                    ("blocked_count", "BIGINT", "NOT NULL", "Количество заблокированных событий."),
                    ("risk_score_sum", "NUMERIC(18,2)", "NOT NULL", "Сумма оценок риска."),
                    ("risk_score_samples", "INTEGER", "NOT NULL", "Количество событий с оценкой риска."),
                    ("last_event_at", "TIMESTAMPTZ", "NOT NULL", "Дата последнего события."),
                ],
            },
            {
                "name": "anomaly_event_rollups",
                "purpose": "Хранит агрегаты аномалий по дате, компьютеру и типу отклонения.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код агрегата."),
                    ("bucket_date", "DATE", "UNIQUE PART", "Дата агрегирования."),
                    ("computer_id", "INTEGER", "UNIQUE PART, NOT NULL", "Код компьютера."),
                    ("anomaly_type", "VARCHAR(100)", "UNIQUE PART, NOT NULL", "Тип аномалии."),
                    ("total_count", "BIGINT", "NOT NULL", "Количество аномалий."),
                    ("high_priority_count", "BIGINT", "NOT NULL", "Количество высокоприоритетных аномалий."),
                    ("last_event_at", "TIMESTAMPTZ", "NOT NULL", "Дата последнего события."),
                ],
            },
        ],
    },
    {
        "name": "База данных NotificationService",
        "physical": "notifications",
        "purpose": (
            "База данных сервиса уведомлений хранит сформированные уведомления, шаблоны сообщений, "
            "признаки доставки и служебные записи для повторной обработки. Она отвечает за информирование "
            "администраторов о значимых событиях и аномалиях."
        ),
        "tables": [
            {
                "name": "notifications",
                "purpose": "Хранит уведомления, сформированные системой.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код уведомления."),
                    ("user_id", "INTEGER", "", "Код получателя."),
                    ("type", "VARCHAR(50)", "", "Тип уведомления."),
                    ("title", "VARCHAR(255)", "", "Заголовок уведомления."),
                    ("message", "TEXT", "", "Текст уведомления."),
                    ("is_read", "BOOLEAN", "", "Признак прочтения."),
                    ("sent_at", "TIMESTAMP", "", "Дата отправки."),
                    ("recipient_email", "VARCHAR(320)", "", "Email получателя."),
                    ("channel", "VARCHAR(20)", "", "Канал доставки."),
                    ("delivery_status", "VARCHAR(32)", "NOT NULL", "Статус доставки."),
                    ("delivery_attempts", "INTEGER", "NOT NULL", "Количество попыток доставки."),
                    ("max_delivery_attempts", "INTEGER", "NOT NULL", "Максимальное число попыток."),
                    ("last_delivery_error", "TEXT", "", "Последняя ошибка доставки."),
                    ("next_retry_at", "TIMESTAMPTZ", "", "Дата следующей попытки."),
                    ("delivered_at", "TIMESTAMPTZ", "", "Дата успешной доставки."),
                ],
            },
            {
                "name": "notification_templates",
                "purpose": "Хранит шаблоны уведомлений.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код шаблона."),
                    ("type", "VARCHAR(50)", "UNIQUE, NOT NULL", "Тип уведомления."),
                    ("subject", "VARCHAR(255)", "", "Тема сообщения."),
                    ("body_template", "TEXT", "", "Шаблон тела сообщения."),
                ],
            },
            {
                "name": "processed_event_inbox",
                "purpose": "Фиксирует обработанные события для предотвращения повторной генерации уведомлений.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код inbox-записи."),
                    ("consumer", "VARCHAR(128)", "NOT NULL", "Название потребителя."),
                    ("event_key", "VARCHAR(256)", "NOT NULL", "Ключ события."),
                    ("message_id", "VARCHAR(128)", "", "Идентификатор сообщения."),
                    ("processed_at", "TIMESTAMP", "NOT NULL", "Дата обработки."),
                ],
            },
            {
                "name": "notification_delivery_dlq",
                "purpose": "Хранит уведомления, которые не удалось доставить после повторных попыток.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код записи DLQ."),
                    ("notification_id", "INTEGER", "UNIQUE, NOT NULL", "Код уведомления."),
                    ("channel", "VARCHAR(20)", "NOT NULL", "Канал доставки."),
                    ("recipient_email", "VARCHAR(320)", "", "Email получателя."),
                    ("attempts", "INTEGER", "NOT NULL", "Количество попыток."),
                    ("reason", "TEXT", "NOT NULL", "Причина ошибки."),
                    ("failed_at", "TIMESTAMPTZ", "NOT NULL", "Дата фиксации ошибки."),
                ],
            },
        ],
    },
    {
        "name": "База данных ReportService",
        "physical": "reports",
        "purpose": (
            "База данных сервиса отчетов хранит отчетные проекции и пользовательскую статистику. "
            "Она позволяет быстро формировать ежедневные отчеты, агрегаты по аномалиям и показатели "
            "активности без обращения к полному журналу событий."
        ),
        "tables": [
            {
                "name": "daily_reports",
                "purpose": "Хранит ежедневные агрегированные отчеты по компьютерам.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код отчета."),
                    ("report_date", "DATE", "UNIQUE PART, NOT NULL", "Дата отчета."),
                    ("computer_id", "INTEGER", "UNIQUE PART", "Код компьютера."),
                    ("user_id", "INTEGER", "", "Код пользователя."),
                    ("total_activities", "BIGINT", "NOT NULL", "Общее количество событий."),
                    ("blocked_actions", "BIGINT", "NOT NULL", "Количество заблокированных действий."),
                    ("avg_risk_score", "NUMERIC(5,2)", "", "Средняя оценка риска."),
                    ("anomaly_count", "BIGINT", "NOT NULL", "Количество аномалий."),
                    ("risk_score_samples", "INTEGER", "NOT NULL", "Количество событий с оценкой риска."),
                    ("created_at", "TIMESTAMP", "", "Дата создания отчета."),
                ],
            },
            {
                "name": "user_stats",
                "purpose": "Хранит статистику пользователя за заданный период.",
                "columns": [
                    ("id", "SERIAL", "PK", "Код записи статистики."),
                    ("user_id", "INTEGER", "", "Код пользователя."),
                    ("period_start", "TIMESTAMP", "NOT NULL", "Начало периода."),
                    ("period_end", "TIMESTAMP", "NOT NULL", "Конец периода."),
                    ("total_time_ms", "BIGINT", "", "Общее активное время."),
                    ("risky_sites", "JSONB", "", "Список рискованных сайтов."),
                    ("violations", "INTEGER", "", "Количество нарушений."),
                    ("created_at", "TIMESTAMP", "", "Дата создания записи."),
                ],
            },
            {
                "name": "processed_event_inbox",
                "purpose": "Фиксирует обработанные события RabbitMQ для корректного построения отчетов.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код inbox-записи."),
                    ("consumer", "VARCHAR(128)", "NOT NULL", "Название потребителя."),
                    ("event_key", "VARCHAR(256)", "NOT NULL", "Ключ события."),
                    ("message_id", "VARCHAR(128)", "", "Идентификатор сообщения."),
                    ("processed_at", "TIMESTAMPTZ", "NOT NULL", "Дата обработки."),
                ],
            },
            {
                "name": "report_daily_anomaly_rollups",
                "purpose": "Хранит ежедневные агрегаты аномалий по компьютерам.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код агрегата."),
                    ("bucket_date", "DATE", "UNIQUE PART, NOT NULL", "Дата агрегирования."),
                    ("computer_id", "INTEGER", "UNIQUE PART, NOT NULL", "Код компьютера."),
                    ("anomaly_type", "VARCHAR(100)", "UNIQUE PART, NOT NULL", "Тип аномалии."),
                    ("total_count", "BIGINT", "NOT NULL", "Количество аномалий."),
                    ("last_event_at", "TIMESTAMPTZ", "NOT NULL", "Дата последнего события."),
                ],
            },
        ],
    },
    {
        "name": "Runtime-база API Gateway",
        "physical": "gateway runtime",
        "purpose": (
            "Runtime-хранилище API Gateway используется для настроек приложения, правил оповещений, "
            "ролевых разрешений и аудита административных операций. Оно поддерживает работу единой "
            "точки входа в систему и разграничение доступа."
        ),
        "tables": [
            {
                "name": "app_settings_documents",
                "purpose": "Хранит настройки приложения в формате JSON.",
                "columns": [
                    ("id", "INTEGER", "PK", "Код документа настроек."),
                    ("payload_json", "TEXT", "NOT NULL", "Содержимое настроек."),
                    ("updated_at", "TIMESTAMP", "NOT NULL", "Дата обновления."),
                ],
            },
            {
                "name": "alert_rules",
                "purpose": "Хранит правила оповещения по метрикам и аномалиям.",
                "columns": [
                    ("id", "UUID", "PK", "Код правила."),
                    ("name", "VARCHAR(255)", "NOT NULL", "Название правила."),
                    ("enabled", "BOOLEAN", "NOT NULL", "Признак активности правила."),
                    ("severity", "VARCHAR(32)", "NOT NULL", "Уровень важности."),
                    ("metric", "VARCHAR(64)", "NOT NULL", "Контролируемая метрика."),
                    ("operator", "VARCHAR(16)", "NOT NULL", "Оператор сравнения."),
                    ("threshold", "NUMERIC(18,4)", "NOT NULL", "Пороговое значение."),
                    ("window_minutes", "INTEGER", "NOT NULL", "Окно анализа в минутах."),
                    ("activity_type", "VARCHAR(64)", "", "Фильтр по типу активности."),
                    ("user_id", "INTEGER", "", "Фильтр по пользователю."),
                    ("computer_id", "INTEGER", "", "Фильтр по компьютеру."),
                    ("notify_in_app", "BOOLEAN", "NOT NULL", "Отправка уведомления в интерфейсе."),
                    ("notify_email", "BOOLEAN", "NOT NULL", "Отправка уведомления по email."),
                    ("cooldown_minutes", "INTEGER", "NOT NULL", "Период подавления повторных уведомлений."),
                    ("created_at", "TIMESTAMP", "NOT NULL", "Дата создания."),
                    ("updated_at", "TIMESTAMP", "NOT NULL", "Дата обновления."),
                ],
            },
            {
                "name": "admin_audit_events",
                "purpose": "Хранит журнал административных действий.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код записи аудита."),
                    ("action", "VARCHAR(128)", "NOT NULL", "Выполненное действие."),
                    ("actor", "VARCHAR(128)", "NOT NULL", "Исполнитель действия."),
                    ("target_type", "VARCHAR(64)", "NOT NULL", "Тип целевого объекта."),
                    ("target_id", "VARCHAR(128)", "NOT NULL", "Идентификатор целевого объекта."),
                    ("success", "BOOLEAN", "NOT NULL", "Признак успешного выполнения."),
                    ("status_code", "INTEGER", "", "HTTP-статус результата."),
                    ("details_json", "TEXT", "NOT NULL", "Детали действия."),
                    ("created_at", "TIMESTAMP", "NOT NULL", "Дата события аудита."),
                ],
            },
            {
                "name": "role_permissions",
                "purpose": "Хранит соответствие ролей и разрешений.",
                "columns": [
                    ("id", "BIGSERIAL", "PK", "Код записи разрешения."),
                    ("role_name", "VARCHAR(128)", "UNIQUE PART, NOT NULL", "Название роли."),
                    ("permission", "VARCHAR(256)", "UNIQUE PART, NOT NULL", "Разрешение."),
                    ("created_at", "TIMESTAMP", "NOT NULL", "Дата создания."),
                    ("updated_at", "TIMESTAMP", "NOT NULL", "Дата обновления."),
                ],
            },
        ],
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def apply_text_style(paragraph, size=9, bold=False, color=None):
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


SUMMARY_PURPOSES = {
    "auth": "Аутентификация, роли и пользовательские сессии.",
    "users": "Профили пользователей и рабочие станции.",
    "activities": "События активности, аномалии, архив и outbox.",
    "agents": "Endpoint-агенты, политики, команды и синхронизация.",
    "metrics": "Правила контроля и агрегированные метрики.",
    "notifications": "Уведомления, шаблоны и статусы доставки.",
    "reports": "Отчетные проекции и пользовательская статистика.",
    "gateway runtime": "Настройки Gateway, RBAC, alert rules и аудит.",
}


def add_table(document, headers, rows, widths, left_cols=(0, 3)):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_text_style(p, size=9, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in left_cols else WD_ALIGN_PARAGRAPH.CENTER
                apply_text_style(p, size=8.5)
    set_table_widths(table, widths)
    return table


def add_paragraph(document, text, style=None):
    p = document.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    return p


def lower_first(text):
    return text[:1].lower() + text[1:] if text else text


def add_narrative_after_table(document, table_name, purpose, db_name=None):
    if db_name:
        text = (
            f"После выделения отношений видно, что {db_name} охватывает отдельный участок "
            "предметной области и не смешивает свои данные с другими подсистемами. Такое "
            "разделение упрощает развитие схемы и делает связи между компонентами более прозрачными."
        )
    else:
        text = (
            f"Таким образом, таблица «{table_name}» {lower_first(purpose)} "
            "Перечисленные поля позволяют однозначно идентифицировать запись, хранить основные "
            "характеристики объекта и использовать эти данные в дальнейшей обработке системы."
        )
    add_paragraph(document, text)


def setup_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_doc():
    document = Document()
    setup_styles(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Описание баз данных и таблиц информационной системы")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Система учета и анализа компьютерной активности пользователей")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_paragraph(
        document,
        "Документ содержит краткое назначение каждой базы данных проекта, перечень отношений "
        "и описание столбцов таблиц с указанием типов данных, ключей и роли каждого поля.",
    )

    document.add_heading("Сводная характеристика баз данных", level=1)
    add_paragraph(
        document,
        "Перед детальным рассмотрением таблиц целесообразно сначала показать, какие базы данных "
        "используются в системе и за какую часть функциональности отвечает каждая из них. Такая "
        "сводка помогает связать физическую структуру хранения с компонентной архитектурой проекта.",
    )
    summary_rows = []
    for db in DATABASES:
        summary_rows.append(
            (
                db["name"],
                db["physical"],
                SUMMARY_PURPOSES[db["physical"]],
                ", ".join(table["name"] for table in db["tables"]),
            )
        )
    add_table(
        document,
        ["База данных", "Физическое имя", "Назначение", "Основные таблицы"],
        summary_rows,
        [Inches(1.55), Inches(1.0), Inches(2.05), Inches(1.90)],
        left_cols=(0, 2, 3),
    )
    add_paragraph(
        document,
        "Из сводной таблицы видно, что каждая база данных обслуживает отдельную подсистему. "
        "Далее каждая из них рассматривается подробнее: сначала приводятся ее отношения в виде "
        "кортежей, затем раскрывается структура каждой таблицы.",
    )

    document.add_heading("Описание таблиц", level=1)
    add_paragraph(
        document,
        "В данном разделе описание построено последовательно: от назначения базы данных к ее "
        "отношениям и далее к составу столбцов. Такой порядок позволяет проследить, как данные "
        "переходят от учетных записей и рабочих станций к событиям активности, аномалиям, "
        "уведомлениям и отчетам.",
    )
    for db_index, db in enumerate(DATABASES, start=1):
        if db_index > 1:
            document.add_page_break()
        document.add_heading(f"{db_index}. {db['name']} ({db['physical']})", level=2)
        add_paragraph(document, db["purpose"])
        add_paragraph(
            document,
            "Для данной базы данных можно выделить следующие отношения. Кортеж отношения показывает "
            "имя таблицы и набор атрибутов, которые описывают соответствующий объект предметной области.",
        )

        tuple_rows = []
        for table in db["tables"]:
            fields = ", ".join(col[0] for col in table["columns"])
            tuple_rows.append((table["name"], f"{table['name']}({fields})"))
        add_table(
            document,
            ["Отношение", "Кортеж отношения"],
            tuple_rows,
            [Inches(1.45), Inches(5.05)],
            left_cols=(0, 1),
        )
        add_narrative_after_table(document, None, None, db["name"])

        for table in db["tables"]:
            document.add_heading(f"Таблица {table['name']}", level=3)
            add_paragraph(
                document,
                f"Таблица «{table['name']}» {lower_first(table['purpose'])} "
                "Ниже приведены ее столбцы, типы данных, ключевые ограничения и назначение каждого поля.",
            )
            add_table(
                document,
                ["Столбец", "Тип данных", "Ключ", "Назначение"],
                table["columns"],
                [Inches(1.45), Inches(1.35), Inches(0.95), Inches(2.75)],
                left_cols=(0, 3),
            )
            add_narrative_after_table(document, table["name"], table["purpose"])

    document.add_heading("Вывод", level=1)
    add_paragraph(
        document,
        "Представленная структура баз данных обеспечивает полный цикл работы системы: "
        "аутентификацию пользователей, хранение профилей и рабочих станций, сбор событий активности, "
        "выявление аномалий, управление endpoint-агентами, построение метрик, формирование уведомлений "
        "и подготовку отчетных данных. Разделение таблиц по сервисам соответствует микросервисной "
        "архитектуре проекта и повышает сопровождаемость информационной системы.",
    )

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Описание таблиц базы данных")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
